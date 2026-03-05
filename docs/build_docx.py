# -*- coding: utf-8 -*-
"""Parse MMFE technical paper markdown and generate filtered Word document.

Reads content from technical_paper_fusion_engine.md.
Filters out: agent system refs, version tags, MGIM refs.
"""
import re
import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# --- Filtering config (Unicode escapes for Chinese) ---

# Heading text substrings that trigger section-level skip
SKIP_HEADINGS = [
    "\u4e0e\u73b0\u6709\u5de5\u4f5c\u7684\u5173\u7cfb",  # 与现有工作的关系
    "\u7cfb\u7edf\u96c6\u6210",                            # 系统集成
    "\u53c2\u8003\u5b9e\u73b0",                            # 参考实现
    "\u6a21\u5757\u4f9d\u8d56\u5173\u7cfb",                # 模块依赖关系
    "\u5df2\u4fee\u590d\u7684\u5c40\u9650\u6027",          # 已修复的局限性
    "\u5df2\u4fee\u590d",                                  # 已修复 (catches sub-headings)
]

# Regex patterns to remove from text
CLEAN_PATTERNS = [
    (r"\s*\(v5\.\d[^)]*\)", ""),                          # (v5.6), (v5.6 新增) — half-width
    (r"\s*\uff08v5\.\d[^\uff09]*\uff09", ""),             # （v5.6 新增）— full-width
    (r"\s*\(v6\.\d[^)]*\)", ""),                          # (v6.0) — half-width
    (r"\s*\uff08v6\.\d[^\uff09]*\uff09", ""),             # （v6.0）— full-width
    (r"\s*\(v7\.\d[^)]*\)", ""),                          # (v7.0), (v7.0新增) — half-width
    (r"\s*\uff08v7\.\d[^\uff09]*\uff09", ""),             # （v7.0 新增）— full-width
    (r"v5\.\d+\s*\u7248\u672c", ""),                      # v5.6 版本
    (r"v6\.\d+\s*\u7248\u672c", ""),                      # v6.0 版本
    (r"v7\.\d+\s*\u7248\u672c", ""),                      # v7.0 版本
    (r"\u501f\u9274[^,\u3002]*\u601d\u60f3\uff0c?", ""),  # 借鉴...思想，
    (r"MGIM\uff08[^\uff09]*\uff09\u7684", ""),            # MGIM（...）的
    (r"MGIM\uff08[^\uff09]*\uff09", ""),                  # MGIM（...）
    (r"MGIM", ""),                                         # MGIM standalone
    (r"\u5e76\u5df2\u96c6\u6210\u5230[^.。]*[.。]", ""),   # 并已集成到...。
    (r"~~([^~]+)~~", r"\1"),                               # ~~strikethrough~~ → text
    (r"\s*v5\.\d+\s*", " "),                               # standalone v5.5 / v5.6
    (r"\s*v6\.\d+\s*", " "),                               # standalone v6.0
    (r"\s*v7\.\d+\s*", " "),                               # standalone v7.0
    (r",?\s*\u542bv5\.\d[^)）]*", ""),                     # 含v5.6增强
    (r",?\s*\u542bv6\.\d[^)）]*", ""),                     # 含v6.0增强
    (r",?\s*\u542bv7\.\d[^)）]*", ""),                     # 含v7.0增强
    (r"\u4ecev5\.\d\u7684\d+\u7ec4\u6269\u5c55", ""),     # 从v5.5的6组扩展
    (r",?\s*\u65b0\u589e\uff09", "\uff09"),               # ， 新增） → ）
    (r",?\s*\u65b0\u589e\)", ")"),                         # ， 新增) → )
    (r"\u5df2\u4fee\u590d\uff09", "\uff09"),              # 已修复） → ）
    (r"\u5df2\u4fee\u590d\)", ")"),                        # 已修复) → )
    (r"for GIS Agent Systems:\s*", ""),                    # English subtitle agent ref
    (r"\u667a\u80fd\u4f53\u7cfb\u7edf\u7684", ""),        # 智能体系统的
    (r"\u667a\u80fd\u4f53\u5e73\u53f0\u4e2d", ""),        # 智能体平台中
    (r"\u667a\u80fd\u4f53\u6846\u67b6\u4e2d", ""),        # 智能体框架中
    (r"GIS\u667a\u80fd\u4f53", "GIS"),                    # GIS智能体 → GIS
    (r"\u9762\u5411GIS\u667a\u80fd\u4f53\u7684", ""),     # 面向GIS智能体的
    (r"\u667a\u80fd\u4f53\u7cfb\u7edf", "\u81ea\u52a8\u5316\u7cfb\u7edf"),  # 智能体系统 → 自动化系统
    (r"\u667a\u80fd\u4f53", "\u81ea\u52a8\u5316\u5f15\u64ce"),              # 智能体 → 自动化引擎
    (r"\uff0c\s*\uff09", "\uff09"),                       # ，） → ）  (trailing comma before full-width paren)
    (r",\s*\)", ")"),                                      # , ) → )  (trailing comma before half-width paren)
]

# Lines containing these → skip entirely
SKIP_LINE_KW = [
    "Masked Geographical Information Model",
    "BaseToolset",
    "FunctionTool",
    "google.adk",
    "fusion_tools.py",
    "fusion_tools",
    "knowledge_graph_tools.py",
    "prompts/general.yaml",
    "agent.py",
    "comparison_MMFE",
    "RELEASE_NOTES",
    "AnalysisToolset",
    "FusionToolset",
    "KnowledgeGraphToolset",
    "drl_engine.py",
    "\u5de5\u5177\u96c6\u5c01\u88c5",                     # 工具集封装
    "\u667a\u80fd\u4f53\u96c6\u6210",                     # 智能体集成
    "\u5de5\u5177\u51fd\u6570",                            # 工具函数
    "\u5de5\u5177\u96c6",                                  # 工具集
    "REST API",
    "Toolset",
    "\u65b0\u589e\u6d4b\u8bd5",                           # 新增测试
    "\u7d2f\u8ba1\u6d4b\u8bd5",                           # 累计测试
    "\u5168\u91cf\u56de\u5f52",                            # 全量回归
    "pre-existing failures",
]

# Section number remapping after removing section 4
SEC_REMAP = {"5": "4", "6": "5", "7": "6", "8": "7", "9": "8"}


def clean_text(text):
    """Apply all cleaning patterns to text."""
    for pat, repl in CLEAN_PATTERNS:
        text = re.sub(pat, repl, text)
    # Clean double spaces left by removals
    text = re.sub(r"  +", " ", text)
    text = text.strip()
    return text


def should_skip_line(text):
    """Check if a line should be skipped entirely."""
    for kw in SKIP_LINE_KW:
        if kw in text:
            return True
    return False


def remap_heading_num(text):
    """Remap section numbers after removing section 4.

    Uses a single match to avoid double-mapping (e.g. 6→5→4 bug).
    Handles both "5.1 ..." and "5 ..." patterns.
    """
    m = re.match(r"^(\d+)", text)
    if m:
        num = m.group(1)
        if num in SEC_REMAP:
            text = SEC_REMAP[num] + text[len(num):]
    return text


def parse_markdown(lines):
    """Parse markdown lines into structured elements."""
    elements = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip()

        # Empty or horizontal rule
        if not line or line == "---":
            i += 1
            continue

        # Heading
        if line.startswith("#"):
            hashes = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            elements.append(("heading", min(hashes, 3), text))
            i += 1
            continue

        # Code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            if i < n:
                i += 1  # skip closing ```
            elements.append(("code", "\n".join(code_lines)))
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            tlines = []
            while i < n and lines[i].rstrip().strip().startswith("|"):
                tlines.append(lines[i].rstrip())
                i += 1
            if len(tlines) >= 2:
                hdrs = [c.strip() for c in tlines[0].split("|")[1:-1]]
                rows = []
                for tl in tlines[2:]:
                    cells = [c.strip() for c in tl.split("|")[1:-1]]
                    if len(cells) == len(hdrs):
                        rows.append(cells)
                elements.append(("table", hdrs, rows))
            continue

        # Bold standalone line
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            elements.append(("bold", line[2:-2]))
            i += 1
            continue

        # Bullet list
        if line.startswith("- "):
            elements.append(("bullet", line[2:]))
            i += 1
            continue

        # Numbered item
        if re.match(r"^\d+\.\s", line):
            elements.append(("numbered", line))
            i += 1
            continue

        # Regular paragraph
        elements.append(("paragraph", line))
        i += 1

    return elements


def filter_elements(elements):
    """Filter out unwanted sections and clean text."""
    filtered = []
    skip_level = 0  # heading level that triggered skip (0 = not skipping)

    for elem in elements:
        etype = elem[0]

        # Check if we should stop skipping
        if skip_level > 0 and etype == "heading":
            level = elem[1]
            if level <= skip_level:
                skip_level = 0  # stop skipping

        # Check if this heading should trigger skip
        if etype == "heading":
            text = elem[2]
            for sh in SKIP_HEADINGS:
                if sh in text:
                    skip_level = elem[1]
                    break
            if skip_level > 0:
                continue

        # Skip if still in skipped section
        if skip_level > 0:
            continue

        # Check line-level skip
        text_content = elem[-1] if isinstance(elem[-1], str) else ""
        if text_content and should_skip_line(text_content):
            continue

        # Clean text content
        if etype == "heading":
            level, text = elem[1], elem[2]
            text = clean_text(text)
            text = remap_heading_num(text)
            # Clean title prefix
            title_prefix = "\u9762\u5411GIS\u667a\u80fd\u4f53\u7684"  # 面向GIS智能体的
            if text.startswith(title_prefix):
                text = text[len(title_prefix):]
            # Skip empty headings after cleaning
            stripped = re.sub(r"[\d.\s]+", "", text).strip()
            if not stripped:
                continue
            filtered.append(("heading", level, text))
        elif etype in ("paragraph", "numbered", "bold", "bullet"):
            text = clean_text(elem[-1])
            if text:
                filtered.append((etype, text))
        elif etype == "code":
            code_text = elem[1]
            # Skip code blocks that are about agent/ADK integration
            if "ADK" in code_text or "BaseToolset" in code_text:
                continue
            code_text = clean_text(code_text)
            if code_text.strip():
                filtered.append(("code", code_text))
        elif etype == "table":
            hdrs, rows = elem[1], elem[2]
            # Clean table cells
            hdrs = [clean_text(h) for h in hdrs]
            # Remove version-related columns (e.g. "版本引入", "实现状态", "版本")
            ver_cols = set()
            for ci, h in enumerate(hdrs):
                if re.search(r"v5\.\d|" + "\u7248\u672c", h):  # 版本
                    ver_cols.add(ci)
            if ver_cols:
                hdrs = [h for ci, h in enumerate(hdrs) if ci not in ver_cols]
            new_rows = []
            for row in rows:
                cleaned_row = [clean_text(c) for c in row]
                # Skip rows mentioning agent/ADK
                row_text = " ".join(cleaned_row)
                if should_skip_line(row_text):
                    continue
                # Remove version columns from row
                if ver_cols:
                    cleaned_row = [c for ci, c in enumerate(cleaned_row) if ci not in ver_cols]
                # Skip rows that still have v5.x references
                row_text2 = " ".join(cleaned_row)
                if re.search(r"v5\.\d", row_text2):
                    continue
                new_rows.append(cleaned_row)
            if new_rows and hdrs:
                filtered.append(("table", hdrs, new_rows))

    return filtered


def renumber_subsections(elements):
    """Fix sub-section numbering gaps after filtering.

    Renumbers X.Y headings sequentially within each major section X.
    For X.Y.Z sub-sub-sections, renumbers Z within each X.Y parent.
    """
    # First pass: collect X.Y headings (not X.Y.Z) per major section
    minor_remap = {}  # {major: {old_minor_str: new_minor_int}}
    for elem in elements:
        if elem[0] == "heading":
            # Match X.Y but NOT X.Y.Z
            m = re.match(r"^(\d+)\.(\d+)(?!\.\d)", elem[2])
            if m:
                major, minor = m.group(1), m.group(2)
                if major not in minor_remap:
                    minor_remap[major] = {}
                if minor not in minor_remap[major]:
                    minor_remap[major][minor] = len(minor_remap[major]) + 1

    result = []
    for elem in elements:
        if elem[0] == "heading":
            text = elem[2]
            # Handle X.Y.Z sub-sub-sections (leave Z as-is, remap Y)
            m3 = re.match(r"^(\d+)\.(\d+)\.(\d+)(.*)", text)
            if m3:
                major, minor, sub, rest = m3.group(1), m3.group(2), m3.group(3), m3.group(4)
                if major in minor_remap and minor in minor_remap[major]:
                    new_minor = minor_remap[major][minor]
                    text = "%s.%d.%s%s" % (major, new_minor, sub, rest)
                result.append(("heading", elem[1], text))
                continue
            # Handle X.Y sections
            m2 = re.match(r"^(\d+)\.(\d+)(.*)", text)
            if m2:
                major, minor, rest = m2.group(1), m2.group(2), m2.group(3)
                if major in minor_remap and minor in minor_remap[major]:
                    new_minor = minor_remap[major][minor]
                    text = "%s.%d%s" % (major, new_minor, rest)
                result.append(("heading", elem[1], text))
                continue
        result.append(elem)
    return result


def build_docx(elements):
    """Build Word document from parsed elements."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.17)

    sty = doc.styles["Normal"]
    sty.font.name = "Times New Roman"
    sty.font.size = Pt(10.5)
    sty.paragraph_format.line_spacing = 1.5
    sty.paragraph_format.space_after = Pt(6)

    for lv in range(1, 4):
        hs = doc.styles["Heading %d" % lv]
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.size = Pt([16, 14, 12][lv - 1])
        if lv == 1:
            hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def sc(cell, text, bold=False, sz=9):
        cell.text = ""
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(sz)
        r.font.name = "Times New Roman"
        if bold:
            r.bold = True

    for elem in elements:
        etype = elem[0]

        if etype == "heading":
            level, text = elem[1], elem[2]
            doc.add_heading(text, level=level)

        elif etype == "paragraph" or etype == "numbered":
            text = elem[1]
            # Handle inline bold **...**
            if "**" in text:
                p = doc.add_paragraph()
                parts = re.split(r"\*\*([^*]+)\*\*", text)
                for j, part in enumerate(parts):
                    if not part:
                        continue
                    r = p.add_run(part)
                    if j % 2 == 1:
                        r.bold = True
            else:
                doc.add_paragraph(text)

        elif etype == "bold":
            p = doc.add_paragraph()
            r = p.add_run(elem[1])
            r.bold = True
            r.font.size = Pt(10.5)

        elif etype == "bullet":
            doc.add_paragraph(elem[1], style="List Bullet")

        elif etype == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(elem[1])
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elif etype == "table":
            hdrs, rows = elem[1], elem[2]
            t = doc.add_table(rows=1 + len(rows), cols=len(hdrs))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, h in enumerate(hdrs):
                sc(t.rows[0].cells[ci], h, bold=True)
            for ri, row in enumerate(rows):
                for ci, v in enumerate(row):
                    if ci < len(hdrs):
                        sc(t.rows[ri + 1].cells[ci], v)

    return doc


def main():
    md_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "technical_paper_fusion_engine.md",
    )
    out_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "MMFE_\u6280\u672f\u6587\u6863.docx",
    )

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    print("Parsing markdown (%d lines)..." % len(lines))
    elements = parse_markdown(lines)
    print("Parsed %d elements" % len(elements))

    print("Filtering content...")
    filtered = filter_elements(elements)
    print("Filtered to %d elements" % len(filtered))

    print("Renumbering sub-sections...")
    filtered = renumber_subsections(filtered)

    print("Building Word document...")
    doc = build_docx(filtered)
    doc.save(out_path)
    print("Saved: %s" % out_path)
    print("Paragraphs: %d, Tables: %d" % (len(doc.paragraphs), len(doc.tables)))


if __name__ == "__main__":
    main()
