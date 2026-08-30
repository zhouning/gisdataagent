#!/usr/bin/env python3
"""Generate a formal Chinese DOCX for the GIS Data Agent Cognitive Runtime design."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "GIS_Data_Agent_Cognitive_Runtime_详细设计说明书.md"
TARGET = ROOT / "GIS_Data_Agent_Cognitive_Runtime_详细设计说明书.docx"

TABLE_CAPTIONS = [
    "修订记录",
    "当前实现与目标状态对照",
    "知识分类与权威性",
    "本体层次与运行时消费者",
    "当前本体资产与生产差距",
    "本体关系类别与控制要求",
    "本体生产级技术选型",
    "本体分阶段生产落地",
    "Operational Ontology 核心类型",
    "重型本体控制面与数据面",
    "重型本体技术候选与 ADR 门",
    "工具副作用等级",
    "数据标准治理输出产物",
    "失败恢复策略",
    "记忆类型与存储",
    "自我进化对象与晋级策略",
    "核心逻辑实体职责",
    "第一阶段技术选型",
    "分阶段实施路线",
    "检索验收指标",
    "本体验收指标",
    "主要风险与应对",
    "待确认事项",
]


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, cn: str = "宋体", en: str = "Times New Roman") -> None:
    run.font.name = en
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cn)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, cn: str, en: str, size: float, bold: bool = False,
                   color: str | None = None) -> None:
    style.font.name = en
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cn)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_inline(paragraph, text: str, size: float = 10.5) -> None:
    text = text.replace("<br>", "\n")
    pattern = r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, 9.2, False, "7A2E2E", cn="等线", en="Consolas")
        elif part.startswith("[") and "](" in part:
            label = part[1:part.index("](")]
            run = paragraph.add_run(label)
            set_run_font(run, size, False, "1F5E8C")
            run.underline = True
        else:
            chunks = part.split("\n")
            for idx, chunk in enumerate(chunks):
                if idx:
                    paragraph.add_run().add_break()
                run = paragraph.add_run(chunk)
                set_run_font(run, size)


def add_caption(document: Document, text: str, figure: bool = True) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, 9.5, False, "4D5B66", cn="宋体")
    r.italic = True


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        if not re.fullmatch(r"[\s:\-|]+", raw):
            rows.append([x.strip() for x in raw.split("|")])
        i += 1
    return rows, i


def add_table(document: Document, rows: list[list[str]], caption: str) -> None:
    add_caption(document, caption, figure=False)
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    repeat_table_header(table.rows[0])
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_inline(p, row[ci] if ci < len(row) else "", 8.8 if cols >= 4 else 9.3)
            if ri == 0:
                shade_cell(cell, "D8E8F3")
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string("17324D")
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_cover(document: Document) -> None:
    for _ in range(3):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GIS Data Agent")
    set_run_font(r, 26, True, "17324D", cn="黑体")
    p.paragraph_format.space_after = Pt(8)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cognitive Runtime 详细设计说明书")
    set_run_font(r, 24, True, "2F6B9A", cn="黑体")
    p.paragraph_format.space_after = Pt(16)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向生产级 GIS 智能体的知识、规划、执行、评价与受控自我进化设计")
    set_run_font(r, 13, False, "4D5B66", cn="楷体")
    for _ in range(6):
        document.add_paragraph()
    meta = [
        ("文档版本", "V1.3"),
        ("编制日期", "2026-07-15"),
        ("设计性质", "目标架构与分阶段实施设计"),
        ("首个试点", "数据标准驱动的空间数据治理"),
        ("本体策略", "轻量权威写模型为基线 + 重型平台条件路线"),
        ("自主等级", "受监督自主执行与受控自我进化"),
        ("证据基线", "Git commit 1421e005b227524ce35d537688a140bd2d8d16e7"),
    ]
    table = document.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        set_cell_margins(table.cell(i, 0), top=120, bottom=120)
        set_cell_margins(table.cell(i, 1), top=120, bottom=120)
        table.cell(i, 0).text = ""
        table.cell(i, 1).text = ""
        p1 = table.cell(i, 0).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        set_run_font(r1, 10.5, True, "17324D", cn="黑体")
        p2 = table.cell(i, 1).paragraphs[0]
        r2 = p2.add_run(v)
        set_run_font(r2, 10.5)
        shade_cell(table.cell(i, 0), "D8E8F3")
    document.add_page_break()


def add_toc(document: Document) -> None:
    h = document.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("目录")
    entries = [
        ("修订记录", "3"),
        ("1. 文档目的与范围", "3"),
        ("2. 当前实现评估", "5"),
        ("3. 总体架构设计", "6"),
        ("4. 运行控制设计", "7"),
        ("5. Cognitive Workspace 与状态机", "8"),
        ("6. 知识与证据设计", "10"),
        ("7. 数据标准知识设计", "11"),
        ("8. 规划与执行设计", "25"),
        ("9. 数据标准驱动治理时序", "26"),
        ("10. 评价、恢复与 HITL", "27"),
        ("11. 记忆设计", "28"),
        ("12. 受控自我进化设计", "29"),
        ("13. 核心逻辑数据设计", "30"),
        ("14. 关键接口设计", "32"),
        ("15. 部署与技术选型", "33"),
        ("16. 安全、可观测性与性能", "35"),
        ("17. 分阶段实施路线", "36"),
        ("18. 验收指标", "36"),
        ("19. 风险与应对", "38"),
        ("20. 实施分解与暂缓原则", "39"),
        ("21. 待确认事项", "40"),
        ("22. 追踪说明", "41"),
        ("23. 结论", "41"),
    ]
    table = document.add_table(rows=len(entries), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(14.0)
    table.columns[1].width = Cm(1.2)
    for row, (title, page) in zip(table.rows, entries):
        left, right = row.cells
        set_cell_margins(left, top=20, start=80, bottom=20, end=80)
        set_cell_margins(right, top=20, start=80, bottom=20, end=80)
        left.text = ""
        right.text = ""
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lp.paragraph_format.line_spacing = 0.95
        add_inline(lp, title, 8.8)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_after = Pt(0)
        rp.paragraph_format.line_spacing = 0.95
        add_inline(rp, page, 8.8)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(4)
    r = note.add_run("目录页码对应 V1.3 正式版版式；后续修订时应同步更新目录。")
    set_run_font(r, 9, False, "6B737B", cn="宋体")
    document.add_page_break()


def configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    set_style_font(document.styles["Normal"], "宋体", "Times New Roman", 10.5)
    normal = document.styles["Normal"]
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [
        ("Title", 24, "17324D"),
        ("Heading 1", 15, "17324D"),
        ("Heading 2", 13, "2F6B9A"),
        ("Heading 3", 11.5, "365A73"),
    ]:
        set_style_font(document.styles[name], "黑体", "Arial", size, True, color)
        document.styles[name].paragraph_format.keep_with_next = True
        document.styles[name].paragraph_format.space_before = Pt(10)
        document.styles[name].paragraph_format.space_after = Pt(6)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("GIS Data Agent Cognitive Runtime 详细设计说明书")
    set_run_font(r, 8.5, False, "6B737B")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("第 ")
    set_run_font(r, 9)
    add_field(footer, "PAGE", "1")
    r = footer.add_run(" 页")
    set_run_font(r, 9)
    set_update_fields(document)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure(document)
    add_cover(document)
    add_toc(document)

    i = 0
    table_index = 0
    in_code = False
    code_lines: list[str] = []
    skipped_title = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if not skipped_title and line.startswith("# "):
            skipped_title = True
            i += 1
            continue
        if skipped_title and line.startswith("**") and i < 8:
            i += 1
            continue
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.3)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(6)
                p_pr = p._element.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F4F6")
                p_pr.append(shd)
                r = p.add_run("\n".join(code_lines))
                set_run_font(r, 8.4, False, "263645", cn="等线", en="Consolas")
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        image = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", line)
        if image:
            path = (SOURCE.parent / image.group(2)).resolve()
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(str(path), width=Cm(16.0))
            add_caption(document, image.group(1))
            i += 1
            continue
        if line.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], 13)
            i += 1
            continue
        if line.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            add_inline(p, line[3:], 15)
            i += 1
            continue
        if line.startswith("# "):
            p = document.add_paragraph(style="Title")
            add_inline(p, line[2:], 20)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            caption = TABLE_CAPTIONS[table_index] if table_index < len(TABLE_CAPTIONS) else f"设计数据表 {table_index + 1}"
            table_index += 1
            add_table(document, rows, f"表 {table_index}  {caption}")
            continue
        if line.startswith(">"):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.right_indent = Cm(0.5)
            p_pr = p._element.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF4F8")
            p_pr.append(shd)
            add_inline(p, line.lstrip("> "), 10.5)
            i += 1
            continue
        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, line[2:], 10.5)
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.55)
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, line, 10.5)
            i += 1
            continue
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(5)
        add_inline(p, line, 10.5)
        i += 1

    document.core_properties.title = "GIS Data Agent Cognitive Runtime 详细设计说明书"
    document.core_properties.subject = "GIS 智能体认知运行时、领域本体、知识、工具执行与受控自我进化"
    document.core_properties.author = "GIS Data Agent Architecture"
    document.core_properties.keywords = "GIS Data Agent, Cognitive Runtime, Domain Ontology, RAG, 数据标准, Agent, HITL, 自我进化"
    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    build()
