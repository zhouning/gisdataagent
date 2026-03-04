# -*- coding: utf-8 -*-
"""Read fusion paper markdown and generate Word document."""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def sc(cell, text, bold=False, sz=9):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text)
    r.font.size = Pt(sz)
    r.font.name = 'Times New Roman'
    if bold:
        r.bold = True


def tbl(doc, hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        sc(t.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            sc(t.rows[ri+1].cells[ci], str(v))
    return t


def code(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(txt)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def bh(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)


# Content stored as a dict to keep code clean
C = {}

def load_content():
    """Load all Chinese content from a data file."""
    C['title'] = '\u9762\u5411GIS\u667a\u80fd\u4f53\u7684\u591a\u6a21\u6001\u7a7a\u95f4\u6570\u636e\u667a\u80fd\u878d\u5408\u5f15\u64ce'
    C['subtitle'] = '\u67b6\u6784\u8bbe\u8ba1\u4e0e\u5b9e\u73b0'
    C['abstract_label'] = '\u6458\u8981'
    C['keywords_label'] = '\u5173\u952e\u8bcd'

load_content()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin, sec.right_margin = Cm(3.17), Cm(3.17)

    sty = doc.styles['Normal']
    sty.font.name = 'Times New Roman'
    sty.font.size = Pt(10.5)
    sty.paragraph_format.line_spacing = 1.5
    sty.paragraph_format.space_after = Pt(6)

    for lv in range(1, 4):
        hs = doc.styles[f'Heading {lv}']
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.size = Pt([16, 14, 12][lv-1])
        if lv == 1:
            hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


if __name__ == '__main__':
    # Read content strings from a separate data file
    data_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fusion_paper_content.py')

    # Generate data file with content
    print('Generating Word document...')

    # Import the content module
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from fusion_paper_content import build_full_doc

    doc = build_full_doc()
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MMFE_\u6280\u672f\u6587\u6863.docx')
    doc.save(out)
    print(f'Saved: {out}')
