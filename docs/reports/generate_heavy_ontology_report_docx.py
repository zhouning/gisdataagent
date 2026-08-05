#!/usr/bin/env python3
"""Generate the formal Word edition of the heavy ontology architecture report."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
DESIGN_ROOT = ROOT.parent / "designs" / "gis_data_agent_cognitive_runtime_2026-07-15"
sys.path.insert(0, str(DESIGN_ROOT))

import generate_docx as base  # noqa: E402


SOURCE = ROOT / "gis-data-agent-heavy-ontology-production-architecture-2026-07-15.md"
TARGET = ROOT / "gis-data-agent-heavy-ontology-production-architecture-2026-07-15.docx"

TABLE_CAPTIONS = [
    "轻量受治理本体与重型本体平台边界",
    "形式语义标准及生产边界",
    "RDF 与语义平台候选",
    "重型平台其他组件候选",
    "主要风险、生产后果与控制原则",
]

TOC_ENTRIES = [
    ("1. 执行摘要", "3"),
    ("2. 轻量与重型的严格边界", "3"),
    ("3. 完整逻辑架构", "4"),
    ("4. 六个控制与数据平面", "5"),
    ("5. 形式语义标准与适用范围", "6"),
    ("6. GIS 特定边界", "6"),
    ("7. 多存储一致性", "7"),
    ("8. 查询网关与运行时协议", "7"),
    ("9. 动态安全架构", "8"),
    ("10. Ontology CI/CD 与发布治理", "8"),
    ("11. 生产部署拓扑", "8"),
    ("12. 技术选型", "9"),
    ("13. 团队与运营要求", "10"),
    ("14. 成本、风险与控制", "10"),
    ("15. 分阶段落地路线", "11"),
    ("16. 准入标准与最终建议", "12"),
    ("17. 证据与限制", "12"),
]


def add_cover(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GIS Data Agent")
    base.set_run_font(r, 26, True, "17324D", cn="黑体")
    p.paragraph_format.space_after = Pt(10)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("重型本体生产架构分析与技术选型")
    base.set_run_font(r, 22, True, "2F6B9A", cn="黑体")
    p.paragraph_format.space_after = Pt(18)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Enterprise Semantic + Operational Ontology Platform 条件目标设计")
    base.set_run_font(r, 12.5, False, "4D5B66", cn="楷体")

    for _ in range(7):
        document.add_paragraph()

    meta = [
        ("文档版本", "V1.0"),
        ("编制日期", "2026-07-15"),
        ("文档性质", "企业级重型本体条件目标架构与技术选型"),
        ("当前决策", "保持轻量 Stage 1/2；重型 H0-H7 未启动"),
        ("适用项目", "GIS Data Agent Cognitive Runtime"),
        ("证据状态", "当前实现、目标设计与 needs-owner-input 分开表述"),
    ]
    table = document.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, value) in enumerate(meta):
        left, right = table.cell(i, 0), table.cell(i, 1)
        base.set_cell_margins(left, top=120, bottom=120)
        base.set_cell_margins(right, top=120, bottom=120)
        left.text = ""
        right.text = ""
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(key)
        base.set_run_font(lr, 10.5, True, "17324D", cn="黑体")
        rr = right.paragraphs[0].add_run(value)
        base.set_run_font(rr, 10.5)
        base.shade_cell(left, "D8E8F3")
    document.add_page_break()


def add_toc(document: Document) -> None:
    h = document.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("目录")
    table = document.add_table(rows=len(TOC_ENTRIES), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(14.0)
    table.columns[1].width = Cm(1.2)
    for row, (title, page) in zip(table.rows, TOC_ENTRIES):
        left, right = row.cells
        base.set_cell_margins(left, top=35, start=80, bottom=35, end=80)
        base.set_cell_margins(right, top=35, start=80, bottom=35, end=80)
        left.text = ""
        right.text = ""
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        base.add_inline(lp, title, 9.3)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_after = Pt(0)
        base.add_inline(rp, page, 9.3)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("目录页码对应 V1.0 正式版版式；后续修订时应同步更新。")
    base.set_run_font(r, 9, False, "6B737B")
    document.add_page_break()


def configure(document: Document) -> None:
    base.configure(document)
    header = document.sections[0].header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("GIS Data Agent 重型本体生产架构分析与技术选型")
    base.set_run_font(r, 8.5, False, "6B737B")


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure(document)
    add_cover(document)
    add_toc(document)

    i = 0
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        i += 1

    table_index = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.3)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(6)
                p_pr = p._element.get_or_add_pPr()
                shd = base.OxmlElement("w:shd")
                shd.set(base.qn("w:fill"), "F2F4F6")
                p_pr.append(shd)
                r = p.add_run("\n".join(code_lines))
                base.set_run_font(r, 8.4, False, "263645", cn="等线", en="Consolas")
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        image = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", line)
        if image:
            path = (SOURCE.parent / image.group(2)).resolve()
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(path), width=Cm(16.0))
            base.add_caption(document, image.group(1))
            i += 1
            continue
        if line.startswith("## "):
            p = document.add_paragraph(style="Heading 2")
            base.add_inline(p, line[3:], 13)
            i += 1
            continue
        if line.startswith("# "):
            p = document.add_paragraph(style="Heading 1")
            base.add_inline(p, line[2:], 15)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = base.parse_table(lines, i)
            caption = TABLE_CAPTIONS[table_index]
            table_index += 1
            base.add_table(document, rows, f"表 {table_index}  {caption}")
            continue
        if line.startswith(">"):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            base.add_inline(p, line.lstrip("> "), 10.5)
            i += 1
            continue
        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            base.add_inline(p, line[2:], 10.5)
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.55)
            base.add_inline(p, line, 10.5)
            i += 1
            continue

        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(5)
        base.add_inline(p, line, 10.5)
        i += 1

    document.core_properties.title = "GIS Data Agent 重型本体生产架构分析与技术选型"
    document.core_properties.subject = "Enterprise Semantic and Operational Ontology Platform 条件目标设计"
    document.core_properties.author = "GIS Data Agent Architecture"
    document.core_properties.keywords = "GIS Data Agent, Ontology, RDF, OWL, SHACL, Operational Ontology, Cognitive Runtime"
    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    build()
