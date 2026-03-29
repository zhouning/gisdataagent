#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Part 3: 安全架构"""
from docx import Document

doc = Document('D:\\adk\\设计文档_完整版_Part2.docx')

# 5. 总体安全架构
doc.add_heading('5. 总体安全架构', 1)
doc.add_paragraph("""采用六层防御体系：""")

doc.add_heading('5.1 认证层', 2)
doc.add_paragraph("""
• 密码认证：PBKDF2-HMAC-SHA256（100k迭代）+ 常量时间比较
• JWT Cookie：HTTP-only，防XSS
• OAuth2：支持Google/GitHub（可选）
• 暴力破解防护：5次失败锁定15分钟
• 默认管理员：首次启动自动创建（admin/admin123）""")

doc.add_heading('5.2 授权层', 2)
doc.add_paragraph("""
• RBAC三角色：
  - admin：完全访问权限
  - analyst：分析管线访问权限
  - viewer：只读权限
• 文件沙箱：uploads/{user_id}/ 目录隔离
• ContextVar传播：current_user_id、current_session_id、current_user_role
• RLS策略：25+张表启用行级安全，基于app.current_user""")

doc.add_heading('5.3 输入验证层', 2)
doc.add_paragraph("""
• SQL注入防护：参数化查询，禁止字符串拼接
• Prompt注入检测：24种禁用模式 + 边界标记检测
• SSRF防护：仅HTTPS，私有IP段阻断
• MCP命令注入：白名单（python/node/docker/uvx等）
• 路径遍历：realpath验证 + 目录前缀检查
• Python沙箱：AST验证，禁止exec/eval/__import__""")

doc.add_heading('5.4 执行隔离层', 2)
doc.add_paragraph("""
• 子进程沙箱：30s超时（最大60s）
• 白名单内置函数：25个允许的builtins
• 环境变量清洗：12个敏感变量移除
• 输出截断：100KB限制（stdout/stderr）""")

doc.add_heading('5.5 输出安全层', 2)
doc.add_paragraph("""
• API密钥混淆：自动检测并替换为***
• 幻觉检测：输出验证机制""")

doc.add_heading('5.6 审计监控层', 2)
doc.add_paragraph("""
• 审计事件：30+类型（登录、数据访问、配置变更等）
• 保留期：90天（可配置）
• Admin仪表板：实时监控 + 历史查询""")

doc.save('D:\\adk\\设计文档_完整版_Part3.docx')
print("Part 3 saved")
