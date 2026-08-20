#!/usr/bin/env python3
# ruff: noqa: E501
"""Build the Abu Dhabi authoritative customer data requirements workbook and Word brief."""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_REGISTER = ROOT / (
    "benchmarks/abu_dhabi_stormwater_data_v1/derived/"
    "abu_dhabi_data_request_register_v2.json"
)
DEFAULT_XLSX = ROOT / (
    "docs/customer/abu_dhabi_liveability_site_validation/"
    "abu_dhabi_stormwater_world_model_authoritative_data_requirements.xlsx"
)
DEFAULT_DOCX = ROOT / (
    "docs/customer/abu_dhabi_liveability_site_validation/"
    "abu_dhabi_stormwater_world_model_authoritative_data_requirements.docx"
)

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
ORANGE = "FCE4D6"
GREEN = "E2F0D9"
GREY = "F2F2F2"
RED = "F4CCCC"
WHITE = "FFFFFF"
THIN_GREY = Side(style="thin", color="B7C9D6")


def _load_register(path: Path) -> dict[str, Any]:
    register = json.loads(path.read_text(encoding="utf-8"))
    if register.get("schema") != "gwm.abu_dhabi_stormwater.data_request_register.v2":
        raise ValueError("unexpected_data_request_register_schema")
    if len(register.get("requests", [])) != 12:
        raise ValueError("expected_twelve_data_request_items")
    return register


def _join(values: Any) -> str:
    if isinstance(values, list):
        return "；".join(str(value) for value in values)
    return str(values or "")


def _set_excel_cell(cell, value: Any, *, wrap: bool = True) -> None:
    cell.value = value
    cell.alignment = Alignment(
        horizontal="left",
        vertical="top",
        wrap_text=wrap,
    )


def _style_excel_header(row) -> None:
    for cell in row:
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.font = Font(color=WHITE, bold=True, size=10)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=Side(style="medium", color=WHITE))


def _style_excel_table(ws, header_row: int, first_data_row: int, last_data_row: int) -> None:
    for row in ws.iter_rows(min_row=header_row, max_row=last_data_row):
        for cell in row:
            cell.border = Border(
                left=THIN_GREY,
                right=THIN_GREY,
                top=THIN_GREY,
                bottom=THIN_GREY,
            )
            if cell.row >= first_data_row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    for row_number in range(first_data_row, last_data_row + 1):
        ws.row_dimensions[row_number].height = 74


def _title_block(ws, title: str, subtitle: str, column_count: int) -> None:
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=column_count)
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=column_count)
    ws.cell(1, 1, title)
    ws.cell(2, 1, subtitle)
    ws.cell(1, 1).font = Font(size=16, bold=True, color=BLUE)
    ws.cell(2, 1).font = Font(size=10, italic=True, color="666666")
    ws.cell(1, 1).alignment = Alignment(horizontal="left", vertical="center")
    ws.cell(2, 1).alignment = Alignment(wrap_text=True, vertical="center")
    ws.row_dimensions[1].height = 28
    ws.row_dimensions[2].height = 42


def build_workbook(register: dict[str, Any], output: Path) -> None:
    requests = register["requests"]
    workbook = Workbook()
    requirements = workbook.active
    requirements.title = "权威数据需求清单"
    columns = [
        "序号",
        "优先级",
        "数据包ID",
        "数据域",
        "客户责任方",
        "客户必须提供的权威真实数据",
        "最低交付内容",
        "关键字段",
        "时间要求",
        "空间/精度要求",
        "建议格式",
        "模型用途",
        "验收检查",
        "当前测试数据/公开替代（仅参考）",
        "当前状态",
        "阻断的模型阶段",
        "客户交付状态",
        "客户备注/文件位置",
    ]
    _title_block(
        requirements,
        "阿布扎比城市暴雨内涝世界模型｜客户权威数据需求清单",
        "基本原则：正式模型只接纳客户或主管部门提供的权威真实数据；公开数据和测试数据仅用于当前管线验证、敏感性诊断和缺口说明，不构成生产输入或模型校准证据。",
        len(columns),
    )
    header_row = 4
    for column, value in enumerate(columns, 1):
        requirements.cell(header_row, column, value)
    _style_excel_header(requirements[header_row])
    for index, item in enumerate(requests, 1):
        row = header_row + index
        values = [
            index,
            item["priority"],
            item["request_id"],
            item["domain"],
            item["requested_from"],
            item["required_artifact"],
            item["minimum_delivery"],
            _join(item["required_fields"]),
            item["time_requirement"],
            item["spatial_requirement"],
            _join(register["delivery_rules"]["accepted_container_formats"]),
            _join(item["model_use"]),
            _join(item["acceptance_checks"]),
            item["current_substitute"] + "（仅参考，不得替代客户权威数据）",
            item["status"],
            _join(item["blocks"]),
            "待客户提供",
            "",
        ]
        for column, value in enumerate(values, 1):
            _set_excel_cell(requirements.cell(row, column), value)
        requirements.cell(row, 1).alignment = Alignment(horizontal="center", vertical="top")
        requirements.cell(row, 2).alignment = Alignment(horizontal="center", vertical="top")
        requirements.cell(row, 17).fill = PatternFill("solid", fgColor=ORANGE)
        requirements.cell(row, 15).fill = PatternFill("solid", fgColor=RED)
        if item["priority"] == "P0":
            requirements.cell(row, 2).fill = PatternFill("solid", fgColor=RED)
        elif item["priority"] == "P1":
            requirements.cell(row, 2).fill = PatternFill("solid", fgColor=ORANGE)
        else:
            requirements.cell(row, 2).fill = PatternFill("solid", fgColor=GREEN)
    last_row = header_row + len(requests)
    _style_excel_table(requirements, header_row, header_row + 1, last_row)
    requirements.auto_filter.ref = f"A{header_row}:{get_column_letter(len(columns))}{last_row}"
    requirements.freeze_panes = "A5"
    widths = [6, 9, 32, 20, 28, 34, 42, 48, 32, 34, 24, 34, 36, 42, 32, 34, 18, 30]
    for index, width in enumerate(widths, 1):
        requirements.column_dimensions[get_column_letter(index)].width = width
    requirements.sheet_properties.tabColor = BLUE
    requirements.conditional_formatting.add(
        f"Q{header_row + 1}:Q{last_row}",
        FormulaRule(formula=[f'Q{header_row + 1}="已接收"'], fill=PatternFill("solid", fgColor=GREEN)),
    )
    status_validation = DataValidation(
        type="list",
        formula1='"待客户提供,已接收待验收,验收通过,退回补充"',
        allow_blank=False,
    )
    requirements.add_data_validation(status_validation)
    status_validation.add(f"Q{header_row + 1}:Q{last_row}")

    principles = workbook.create_sheet("交付原则与状态")
    _title_block(
        principles,
        "交付原则与状态定义",
        "该页用于客户、项目经理和模型团队对“权威数据”和“测试替代数据”建立统一口径。",
        4,
    )
    principle_rows = [
        ("原则", "正式模型输入", "必须来自客户或主管部门的权威真实数据，并具备完整元数据、版本和质量证据。", "客户交付"),
        ("原则", "测试/公开数据", "仅用于管线连通性、格式验证、敏感性诊断和候选空间分析；不得用于生产校准、GWM 训练标签或城市级预测声明。", "当前已有"),
        ("原则", "文件接收不等于准入", "收到文件后仍需完成来源、时间、CRS/垂直基准、单位、质量、哈希、拓扑和盲测验收。", "双方验收"),
        ("状态", "待客户提供", "当前尚未收到客户权威数据。", "阻断对应模型阶段"),
        ("状态", "已接收待验收", "已收到文件，但尚未通过字段、空间、时间、质量和一致性检查。", "不得进入生产模型"),
        ("状态", "验收通过", "完成数据质量审查并形成版本化回执。", "可申请下一 gate"),
        ("状态", "退回补充", "字段缺失、单位/基准不清、时间不完整、哈希或拓扑不一致。", "继续阻断"),
    ]
    headers = ["类型", "主题", "说明", "责任/后果"]
    for col, value in enumerate(headers, 1):
        principles.cell(4, col, value)
    _style_excel_header(principles[4])
    for row_index, values in enumerate(principle_rows, 5):
        for col, value in enumerate(values, 1):
            _set_excel_cell(principles.cell(row_index, col), value)
    _style_excel_table(principles, 4, 5, 4 + len(principle_rows))
    principles.column_dimensions["A"].width = 12
    principles.column_dimensions["B"].width = 24
    principles.column_dimensions["C"].width = 90
    principles.column_dimensions["D"].width = 28
    principles.freeze_panes = "A5"
    principles.sheet_properties.tabColor = "70AD47"

    model_map = workbook.create_sheet("模型与数据对应关系")
    _title_block(
        model_map,
        "传统模型与 GWM 的数据分工",
        "GWM 当前定位为传统模型之上的快速状态代理和情景筛选层；传统模型仍承担物理约束、校准、验证和高风险场景回退。",
        6,
    )
    model_headers = ["模型/层", "当前角色", "主要输入", "主要输出", "是否依赖客户权威数据", "当前准入状态"]
    for col, value in enumerate(model_headers, 1):
        model_map.cell(4, col, value)
    _style_excel_header(model_map[4])
    model_rows = [
        ("EPA SWMM 5.2.4", "一维雨水管网和控制基线", "权威管网、产流参数、事件雨量、泵闸、外排边界", "节点水位、管段流量、溢流和控制响应", "是", "未准入；当前仅合成/候选诊断"),
        ("ANUGA 2D", "二维地表浅水和积水扩散", "工程 DTM、道路/建筑障碍、面雨量、边界条件", "水深、流速、淹没范围和退水过程", "是", "未准入；当前仅候选诊断"),
        ("LISFLOOD-FP", "独立二维诊断/交叉检查", "工程地形、降雨、边界和阻力参数", "二维水深和范围对比", "是", "诊断候选"),
        ("GWM", "快速状态模拟、情景筛选、不确定性和分布偏移检测", "已验收观测事件、已准入传统模型滚动、运行操作历史", "快速预测候选、风险区间、候选行动排序", "是", "训练和城市级预测均关闭"),
        ("混合规划器", "将 GWM 候选行动交由传统模型复核", "资产约束、泵闸规则、候选行动、传统模型验证结果", "经物理校验的行动方案", "是", "仅架构合同，未执行"),
    ]
    for row_index, values in enumerate(model_rows, 5):
        for col, value in enumerate(values, 1):
            _set_excel_cell(model_map.cell(row_index, col), value)
    _style_excel_table(model_map, 4, 5, 4 + len(model_rows))
    for col, width in enumerate([22, 34, 60, 48, 22, 30], 1):
        model_map.column_dimensions[get_column_letter(col)].width = width
    model_map.freeze_panes = "A5"
    model_map.sheet_properties.tabColor = "ED7D31"

    template = workbook.create_sheet("客户交付填报模板")
    _title_block(
        template,
        "客户权威数据交付填报模板",
        "每个文件或数据库快照至少填写一行；如果一个数据包包含多个文件，请分别登记并保持同一数据包 ID。",
        16,
    )
    template_headers = [
        "数据包ID",
        "文件/表名",
        "来源责任方",
        "版本/快照ID",
        "有效时间起（UTC）",
        "有效时间止（UTC）",
        "水平CRS",
        "垂直基准/大地水准面",
        "单位",
        "质量标识说明",
        "许可/复用授权",
        "覆盖范围/空间精度",
        "文件格式/数据库表",
        "SHA-256",
        "客户确认人/日期",
        "验收结果/备注",
    ]
    for col, value in enumerate(template_headers, 1):
        template.cell(4, col, value)
    _style_excel_header(template[4])
    for row in range(5, 25):
        for col in range(1, len(template_headers) + 1):
            _set_excel_cell(template.cell(row, col), "")
        template.row_dimensions[row].height = 40
    _style_excel_table(template, 4, 5, 24)
    template.auto_filter.ref = "A4:P24"
    template.freeze_panes = "A5"
    for col, width in enumerate([32, 32, 28, 22, 22, 22, 18, 24, 16, 32, 30, 32, 22, 70, 24, 36], 1):
        template.column_dimensions[get_column_letter(col)].width = width
    template.sheet_properties.tabColor = "A5A5A5"

    workbook.properties.title = "阿布扎比城市暴雨内涝世界模型客户权威数据需求"
    workbook.properties.subject = "客户权威真实数据交付清单与验收模板"
    workbook.properties.creator = "Abu Dhabi Stormwater World Model"
    output.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _doc_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_text(cell, header, bold=True, color=WHITE)
        _set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if len(table.rows) % 2 == 0:
                _set_cell_shading(cells[index], PALE_BLUE)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def _doc_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _doc_paragraph(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        text = text[len(bold_prefix):]
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _doc_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build_docx(register: dict[str, Any], output: Path, workbook_path: Path) -> None:
    requests = register["requests"]
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [("Title", 22, BLUE), ("Heading 1", 16, BLUE), ("Heading 2", 13, "2F5597")]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("阿布扎比城市暴雨内涝世界模型\n客户权威数据需求说明书")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"客户交付版｜{date.today().isoformat()}")
    run.font.color.rgb = RGBColor.from_string("666666")
    document.add_paragraph()

    principle_table = document.add_table(rows=3, cols=2)
    principle_table.style = "Table Grid"
    principle_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    principle_rows = [
        ("正式输入原则", "世界模型的正式生产输入、传统模型校准数据和 GWM 训练/验证数据，必须由客户或主管部门提供权威真实数据。"),
        ("测试数据原则", "当前公开数据、SmartMakani 候选数据、合成数据和公开再分析数据只用于测试数据管线、格式检查、诊断敏感性和缺口识别，不作为生产数据。"),
        ("准入原则", "文件收到不等于模型准入。每项数据必须通过来源、版本、时间、CRS/垂直基准、单位、质量标识、空间覆盖、哈希和一致性验收。"),
    ]
    for row_index, (label, text) in enumerate(principle_rows):
        _set_cell_text(principle_table.cell(row_index, 0), label, bold=True, color=BLUE)
        _set_cell_shading(principle_table.cell(row_index, 0), LIGHT_BLUE)
        _set_cell_text(principle_table.cell(row_index, 1), text)
    document.add_paragraph()

    _doc_heading(document, "1. 交付结论", 1)
    _doc_paragraph(document, "本说明书不是要求客户提供可公开下载的替代数据，而是将阿布扎比城市暴雨内涝世界模型所需的权威真实数据逐项登记，便于客户按数据包交付。当前使用的 CHIRPS、SmartMakani、Copernicus/SRTM、OSM、公开边界和受限遥感目录等，只是测试或参考证据。正式模型不会因为这些测试数据存在而自动打开任何准入 gate。")
    summary_rows = [
        ["需求总数", "12 项"],
        ["P0 必须先提供", "6 项：地形、管网、事件雨量、潮位外边界、泵闸运行、积水观测"],
        ["P1 后续提供", "5 项：空间叠加、Liveability 暴露、土地覆盖入渗、道路障碍、历史事件/设计暴雨"],
        ["P2 增强项", "1 项：维护、堵塞和设施状态历史"],
        ["当前客户权威数据", "尚未收到；所有 12 项均保持待客户提供"],
    ]
    _doc_table(document, ["项目", "当前情况"], summary_rows, [4.0, 22.0])

    _doc_heading(document, "2. 当前技术架构与模型分工", 1)
    _doc_paragraph(document, "技术架构采用“传统物理模型负责可信边界，GWM 负责快速状态代理和情景筛选”的协作方式。传统模型不会被当前 GWM 直接替代。只有在完成客户权威数据验收、传统模型事件校准、跨事件盲测和 GWM 影子验证后，GWM 才能承担部分高成本内循环搜索或快速预报任务。")
    architecture_rows = [
        ["EPA SWMM 5.2.4", "一维雨水管网、产流、节点、管段、泵闸、外排边界和质量守恒基线。", "权威管网和设施关系、事件雨量、泵闸运行、潮位/外排边界。", "传统物理基线；当前未准入。"],
        ["ANUGA 2D", "二维地表浅水、道路低点、建筑/路缘障碍、积水扩散和退水。", "工程 DTM/LiDAR、微地形、面雨量、边界和观测。", "传统物理基线；当前仅候选诊断。"],
        ["LISFLOOD-FP", "独立二维诊断和交叉检查。", "工程地形、降雨、边界和阻力参数。", "诊断候选，不作为唯一依据。"],
        ["GWM", "学习已验收的“强迫—传统模型状态—观测证据”关系，提供快速滚动、不确定性和分布偏移检测。", "已验收观测事件、已准入传统模型滚动、操作历史。", "训练关闭；不得作为唯一物理权威。"],
        ["混合规划器", "由 GWM 先筛选候选行动，再由传统模型复核质量守恒、边界和资产约束。", "资产动作边界、泵闸规则、传统模型复核结果。", "仅架构合同，未执行。"],
    ]
    _doc_table(document, ["模型/层", "职责", "关键客户数据", "当前状态"], architecture_rows, [3.2, 7.0, 8.5, 5.0])

    _doc_heading(document, "3. 客户权威数据需求", 1)
    _doc_paragraph(document, "下表是 Excel 清单的摘要。完整字段、时间分辨率、空间要求和验收检查请以随附 Excel 的“权威数据需求清单”页为准。")
    overview_rows = []
    for item in requests:
        overview_rows.append(
            [
                item["priority"],
                item["request_id"],
                item["requested_from"],
                item["minimum_delivery"],
                _join(item["model_use"]),
                "待客户提供",
            ]
        )
    _doc_table(document, ["优先级", "数据包 ID", "客户责任方", "最低交付内容", "模型用途", "状态"], overview_rows, [1.5, 3.2, 4.0, 7.0, 5.0, 2.0])

    _doc_heading(document, "4. 客户交付格式与元数据要求", 1)
    _doc_paragraph(document, "客户可以按文件、数据库表或版本化快照交付。推荐使用 GeoTIFF/COG、GeoPackage、GeoJSON、CSV、Parquet、NetCDF、HDF5 或 ZIP；数据库交付需要同时提供快照标识和结构说明。")
    for text in [
        "每个数据包必须有明确的 source owner、版本或快照 ID、获取/有效时间、水平 CRS 或垂直基准、单位、质量标识、许可/复用授权和 SHA-256。",
        "所有事件时间统一使用 UTC；如果源系统使用阿布扎比本地时间，必须提供可复现的转换规则和时钟漂移说明。",
        "高程数据必须同时说明水平 CRS、垂直基准、椭球/大地水准面模型、NoData 规则和精度；不能只提供一个没有基准说明的 DEM。",
        "网络数据必须提供确定性的资产 ID、上下游关系、管径/断面、长度、管底高程、节点高程、泵闸/外排口关系和重复/孤立记录处理规则。",
        "观测数据必须保留原始值、质控值、质量标识和不确定性；训练/盲测划分由双方冻结，不能在训练过程中临时调整。",
    ]:
        _doc_bullet(document, text)
    _doc_paragraph(document, f"客户可以直接填写随附 Excel 的“客户交付填报模板”页：{workbook_path.name}。")

    _doc_heading(document, "5. 数据验收与模型准入路径", 1)
    flow_rows = [
        ["1. 接收登记", "记录文件/表名、来源责任方、版本、时间范围、格式和 SHA-256。", "不改变任何模型状态。"],
        ["2. 基础质量验收", "检查 CRS、垂直基准、单位、时间、质量标识、空间覆盖、缺失值和文件完整性。", "不通过则退回补充。"],
        ["3. 工程一致性验收", "检查管网拓扑、设施交叉表、地形覆盖、边界映射、资产 ID 和版本一致性。", "通过后才可申请 K0。"],
        ["4. 传统模型校准", "用至少一个训练事件和一个独立盲测事件校准 SWMM/ANUGA，并核对水量、深度、范围和退水。", "传统模型获得候选准入。"],
        ["5. GWM 影子验证", "用已准入传统模型滚动和观测标签进行训练/验证，检查不确定性、分布偏移和高风险回退。", "GWM 仍需独立 gate。"],
        ["6. 混合规划", "GWM 筛选候选动作，传统模型复核质量守恒、边界和资产动作约束。", "未通过则回退传统模型或不作声明。"],
    ]
    _doc_table(document, ["阶段", "工作内容", "准入影响"], flow_rows, [3.5, 14.5, 5.0])
    _doc_paragraph(document, "当前状态：K0、传统模型准入、GWM 训练、混合规划器和城市级预测声明全部关闭。")

    _doc_heading(document, "6. 当前测试数据说明", 1)
    _doc_paragraph(document, "为避免客户误解，当前测试数据和公开数据的定位如下：")
    test_rows = [
        ["CHIRPS 2024-04-15..17 日降雨", "公开测试代理", "仅诊断敏感性；日尺度、0.05°，不能替代 NCM 站点/雷达事件雨量。"],
        ["SmartMakani 管网、等高线、设施候选", "公开/客户候选数据", "用于候选网络和空间交叉检查；单位、垂直基准、权威关系和工程版本未验收。"],
        ["Copernicus/SRTM 地形", "公开测试代理", "用于管线和算法测试；不能替代工程 DTM/LiDAR 和微地形。"],
        ["OSM 水系、OCHA/HDX 边界", "公开空间上下文", "用于裁剪和空间关联；不是客户权威资产或积水观测。"],
        ["Sentinel-1/GPM 目录和访问探针", "目录级证据", "没有被当作洪水影像或半小时降雨数据；受限资产不进入模型。"],
        ["SWMM/ANUGA/LISFLOOD-FP 合成运行", "合成诊断", "用于验证适配器、质量守恒和接口；不代表阿布扎比事件重建结果。"],
    ]
    _doc_table(document, ["数据", "性质", "使用边界"], test_rows, [5.0, 3.5, 14.5])

    _doc_heading(document, "7. 客户需要优先提供的 P0 数据", 1)
    _doc_paragraph(document, "如果客户暂时无法一次性提供全部数据，建议按以下顺序分批交付：")
    for text in [
        "第一批：工程 DTM/LiDAR 与垂直基准、权威雨水管网及设施关系、2024 年 4 月事件雨量站/雷达 QPE。",
        "第二批：潮位/风暴增水/外排边界、泵闸运行和调蓄日志、2024 年事件积水深度/范围/退水观测。",
        "第三批：共同空间叠加规则、Liveability 暴露口径、道路/缘石/建筑微地形、土地覆盖/入渗参数。",
        "第四批：历史事件库、设计暴雨库、维护堵塞和设施状态历史。",
    ]:
        _doc_bullet(document, text)
    _doc_paragraph(document, "客户数据到达后，项目组会先更新数据需求回执和验收状态，再决定是否推进下一模型 gate；不会因为测试数据可用而跳过客户权威数据验收。")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = "阿布扎比城市暴雨内涝世界模型客户权威数据需求说明书"
    document.core_properties.subject = "客户权威真实数据交付、验收和模型准入说明"
    document.core_properties.author = "Abu Dhabi Stormwater World Model"
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--register", type=Path, default=DEFAULT_REGISTER)
    parser.add_argument("--xlsx", type=Path, default=DEFAULT_XLSX)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    args = parser.parse_args()
    register = _load_register(args.register.resolve())
    xlsx = args.xlsx.resolve()
    docx = args.docx.resolve()
    build_workbook(register, xlsx)
    build_docx(register, docx, xlsx)
    print(json.dumps({"status": "ok", "xlsx": str(xlsx), "docx": str(docx)}))


if __name__ == "__main__":
    main()
