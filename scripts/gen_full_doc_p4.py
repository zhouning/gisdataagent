#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Part 4: 核心模块架构设计"""
from docx import Document

doc = Document('D:\\adk\\设计文档_完整版_Part3.docx')

# 6. 核心模块架构设计
doc.add_heading('6. 核心模块架构设计', 1)

# 6.1 Intent Router模块
doc.add_heading('6.1 Intent Router模块', 2)
doc.add_heading('6.1.1 简介', 3)
doc.add_paragraph("""Intent Router是系统的语义路由核心，负责将用户自然语言请求分类到对应的处理管线。

技术实现：基于Gemini 2.0 Flash模型
支持语言：中文、英文、日文
分类结果：GOVERNANCE（治理）、OPTIMIZATION（优化）、GENERAL（通用）、WORKFLOW（工作流）、AMBIGUOUS（模糊）""")

doc.add_heading('6.1.2 流程图', 3)
doc.add_paragraph("""
用户输入 → 语言检测 → 多模态处理（图片/PDF） → Gemini分类 → 工具类别提取 → 返回意图+工具列表

关键步骤：
1. 语言检测：基于Unicode字符分布（CJK/Latin/Hiragana）
2. 多模态：图片resize至512px，PDF截取前2000字符
3. 分类：Gemini 2.0 Flash，低延迟
4. 工具提取：返回逗号分隔的工具类别（spatial_processing、poi_location等）""")

doc.add_heading('6.1.3 关键类', 3)
doc.add_paragraph("""
主要函数：
• classify_intent(user_message, images, files) → (intent, tool_categories)
• detect_language(text) → 'zh'|'en'|'ja'
• should_decompose(text) → bool（检测多步骤任务）
• generate_analysis_plan(text) → plan_text（生成执行计划）

配置参数：
• ROUTER_MODEL: gemini-2.0-flash（环境变量）
• 图片批次：最多3张
• PDF截断：2000字符""")

# 6.2 Pipeline Engine模块
doc.add_heading('6.2 Pipeline Engine模块', 2)
doc.add_heading('6.2.1 简介', 3)
doc.add_paragraph("""Pipeline Engine负责三条专业管线的编排和执行。

三条管线：
• Optimization Pipeline：数据探索 → 处理 → 分析质量循环 → 可视化 → 总结
• Governance Pipeline：治理探索 → 处理 → 报告循环
• General Pipeline：通用处理 → 可视化 → 总结循环

设计模式：Generator-Critic循环，最多3次迭代""")

doc.add_heading('6.2.2 流程图', 3)
doc.add_paragraph("""
Intent分类 → 选择管线 → 初始化Agent → 执行步骤 → 质量检查 → 循环/输出

Optimization Pipeline详细流程：
ParallelAgent(Exploration ‖ SemanticPreFetch) → DataProcessing → AnalysisQualityLoop(Analysis → QualityChecker → 反馈) → DataVisualization → DataSummary""")

doc.add_heading('6.2.3 关键类', 3)
doc.add_paragraph("""
主要类：
• SequentialAgent：顺序执行多个子Agent
• ParallelAgent：并行执行多个子Agent
• LoopAgent：循环执行直到满足条件（最多3次）
• LlmAgent：单个LLM驱动的Agent

工厂函数：
• _make_planner_explorer()
• _make_planner_processor()
• _make_planner_analyzer()
• _make_planner_visualizer()

说明：ADK要求每个Agent只能有一个父Agent，因此使用工厂函数创建独立实例""")

doc.save('D:\\adk\\设计文档_完整版_Part4.docx')
print("Part 4 saved")
