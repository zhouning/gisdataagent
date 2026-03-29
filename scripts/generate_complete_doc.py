#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Complete Design Document Generator - All Sections in One Script"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading_1(doc, text):
    h = doc.add_heading(text, 1)
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(text, 2)
    return h

def add_heading_3(doc, text):
    h = doc.add_heading(text, 3)
    return h

def add_para(doc, text):
    p = doc.add_paragraph(text)
    return p

def add_table_data(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    return table

print("生成设计文档...")

doc = Document()

# 标题
title = doc.add_heading('时空数据中台产品详细设计V3.0.0.0（Data Agent部分）', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. 概述
add_heading_1(doc, '1. 概述')
add_para(doc, """本文档描述时空数据中台产品v3.0.0.0的系统架构详细设计，用于指导产品开发实现。本产品基于Google Agent Developer Kit (ADK) v1.27构建，是企业级地理信息智能分析平台。

产品定位：通过LLM驱动的语义路由实现数据治理、土地优化和空间智能三大核心能力
版本：v15.8 (2026-03)
规模：96测试文件、2680+用例、202 API、48表、36工具集
特性：多模态融合、自服务扩展、测绘质检、可观测性、多租户隔离""")

# 2. 总体技术架构
add_heading_1(doc, '2. 总体技术架构')
add_para(doc, """系统采用分层微服务架构：

接入层：Chainlit UI Server (8000端口) + REST API Gateway (202端点) + OAuth2
应用层：Intent Router + 3条Pipeline + Dynamic Planner + Workflow Engine
工具层：36 Toolsets + MCP Hub (stdio/SSE/HTTP) + User Tools Engine
数据层：PostgreSQL 16 + PostGIS 3.4 + Redis + 对象存储
AI层：Model Gateway + Context Manager + Prompt Registry + Eval Framework
监控层：Prometheus (25+指标) + 结构化日志 + Alert Engine""")

# 3. 总体组件/服务架构
add_heading_1(doc, '3. 总体组件/服务架构')
add_para(doc, """核心服务组件：""")
add_table_data(doc, ['组件', '端口', '职责'], [
    ['Chainlit Server', '8000', 'Web UI + WebSocket连接'],
    ['Intent Router', '-', 'Gemini 2.0 Flash语义分类'],
    ['Pipeline Orchestrator', '-', '3条专业管线调度'],
    ['MCP Hub', '-', 'Model Context Protocol集成'],
    ['Workflow Engine', '-', 'DAG执行 + Cron调度'],
    ['CV Detection', '8010', '视觉检测服务'],
    ['CAD Parser', '8011', 'CAD/3D解析服务'],
    ['Reference Data', '8012', '参考数据服务'],
])

# 继续...保存第一部分
doc.save('D:\\adk\\设计文档_完整版.docx')
print("✓ 文档已生成")
