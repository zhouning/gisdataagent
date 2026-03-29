#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Part 6: 非功能性设计 + 部署架构"""
from docx import Document

doc = Document('D:\\adk\\设计文档_完整版_Part5.docx')

# 8. 非功能性设计
doc.add_heading('8. 非功能性设计', 1)

doc.add_heading('8.1 性能设计', 2)
doc.add_paragraph("""
• 数据库连接池：pool_size=5, max_overflow=10, pool_recycle=1800s
• 语义层缓存：5分钟TTL，写入失效
• 分块I/O：大文件分块读取（fusion模块）
• PostGIS下推：大数据集在数据库内执行空间操作
• 异步执行：后台任务队列（agent_task_queue）
• 响应时间目标：API < 200ms，管线执行 < 60s""")

doc.add_heading('8.2 安全性设计', 2)
doc.add_paragraph("""
• 认证：PBKDF2 100k迭代 + JWT Cookie
• 授权：RBAC + 文件沙箱 + RLS
• 输入验证：SQL/Prompt/SSRF/路径检查
• 审计：30+事件类型，90天保留
• 加密：传输层TLS，存储层预留列级加密""")

doc.add_heading('8.3 高可用性设计', 2)
doc.add_paragraph("""
• 无状态应用：支持水平扩展
• 数据库：PostgreSQL主从复制 + PITR备份
• 熔断器：5次失败触发，120s冷却
• 健康检查：/health（存活）、/ready（就绪）
• 优雅降级：数据库不可用时使用YAML配置""")

doc.add_heading('8.4 扩展性设计', 2)
doc.add_paragraph("""
• 自定义Skills：用户定义Agent行为
• User Tools：声明式工具模板（http/sql/file/chain）
• MCP集成：外部工具桥接
• 工作流编排：DAG可视化编辑
• 插件化Toolset：BaseToolset继承体系""")

# 9. 部署架构设计
doc.add_heading('9. 部署架构设计', 1)

doc.add_heading('9.1 传统环境部署模式', 2)
doc.add_paragraph("""
单机部署，适用于开发测试环境。

部署方式：Docker Compose
配置文件：docker-compose.yml

服务清单：
• app：主应用（端口8000）
• db：PostgreSQL 16 + PostGIS（端口5433）
• redis：Redis 7（端口6379，可选）

资源需求：
• CPU：4核
• 内存：8GB
• 磁盘：50GB SSD
• 网络：10Mbps""")

doc.add_heading('9.2 云原生环境非高可用部署模式', 2)
doc.add_paragraph("""
Kubernetes单副本部署，适用于预生产环境。

部署方式：kubectl apply -k k8s/
命名空间：gis-agent

资源配置：
• app Pod：requests(250m CPU, 512Mi内存), limits(2 CPU, 2Gi内存)
• db StatefulSet：requests(250m CPU, 512Mi内存), limits(1 CPU, 2Gi内存)
• PVC：uploads 5Gi, pgdata 10Gi

健康检查：
• readiness：GET /ready，30s延迟，10s间隔
• liveness：GET /health，60s延迟，30s间隔""")

doc.add_heading('9.3 云原生环境高可用部署模式', 2)
doc.add_paragraph("""
Kubernetes多副本部署，适用于生产环境。

HPA配置：
• 最小副本：1
• 最大副本：5
• CPU目标：70%
• 内存目标：80%
• 扩容速率：+2 pods/60s
• 缩容速率：-1 pod/60s，稳定期300s

数据库HA：
• PostgreSQL主从复制
• PgBouncer连接池
• 自动故障转移（Patroni）

存储：
• ReadWriteMany PVC（共享uploads）
• 对象存储（OBS/S3/GCS）""")

doc.save('D:\\adk\\设计文档_完整版_Part6.docx')
print("Part 6 saved")
