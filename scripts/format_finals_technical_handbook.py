"""Apply the final A4 layout to the Gemma 4 technical handbook DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _set_east_asia_font(style, name: str) -> None:
    style.font.name = name
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def _heading_level(paragraph) -> int | None:
    style_name = paragraph.style.name
    if style_name == "Heading 1":
        return 1
    if style_name == "Heading 2":
        return 2
    return None


def _heading_bookmark(paragraph) -> str | None:
    for bookmark in paragraph._p.iter(qn("w:bookmarkStart")):
        name = bookmark.get(qn("w:name"))
        if name and not name.startswith("_"):
            return name

    # Pandoc emits heading bookmarks as body-level siblings immediately before
    # the paragraph, while Word may move the same bookmark inside it.
    sibling = paragraph._p.getprevious()
    while sibling is not None and sibling.tag == qn("w:bookmarkStart"):
        name = sibling.get(qn("w:name"))
        if name and not name.startswith("_"):
            return name
        sibling = sibling.getprevious()
    return None


def _ensure_toc_style(doc, name: str, level: int) -> None:
    if name in doc.styles:
        style = doc.styles[name]
    else:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]

    _set_east_asia_font(style, "PingFang SC")
    style.font.size = Pt(9.5 if level == 1 else 8.75)
    style.font.bold = level == 1
    style.font.color.rgb = RGBColor(0x20, 0x2A, 0x35)
    style.paragraph_format.left_indent = Cm(0 if level == 1 else 0.65)
    style.paragraph_format.space_before = Pt(1.5 if level == 1 else 0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = Pt(10.5)


def _add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = paragraph.add_run(text)
    run.font.name = "PingFang SC"
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.underline = False
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), "PingFang SC"
    )
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def _insert_clickable_toc(doc) -> None:
    has_toc = any(
        paragraph.text.strip() == "目录"
        and paragraph.style.name == "Heading 1"
        for paragraph in doc.paragraphs
    )
    if has_toc:
        _ensure_toc_style(doc, "TOC 1", 1)
        _ensure_toc_style(doc, "TOC 2", 2)
        return

    targets = []
    for paragraph in doc.paragraphs:
        level = _heading_level(paragraph)
        if level is None:
            continue
        text = paragraph.text.strip()
        bookmark = _heading_bookmark(paragraph)
        if text and bookmark:
            targets.append((paragraph, level, text, bookmark))

    if not targets:
        raise ValueError("No bookmarked Heading 1/2 paragraphs found for the TOC")

    _ensure_toc_style(doc, "TOC 1", 1)
    _ensure_toc_style(doc, "TOC 2", 2)

    first_heading = targets[0][0]
    toc_heading = first_heading.insert_paragraph_before("目录", style="Heading 1")
    toc_heading.paragraph_format.page_break_before = True
    toc_heading.paragraph_format.space_after = Pt(4)

    note = first_heading.insert_paragraph_before()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.space_after = Pt(8)
    note_run = note.add_run(
        "点击目录条目可跳转至正文；在 Word 中打开“导航窗格”可使用完整章节树。"
    )
    note_run.font.name = "PingFang SC"
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    note_run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), "PingFang SC"
    )

    chapter_starts = [
        index for index, (_, level, _, _) in enumerate(targets) if level == 1
    ]
    split_index = min(
        (index for index in chapter_starts if index > 0),
        key=lambda index: abs(index - len(targets) / 2),
    )

    for index, (_, level, text, bookmark) in enumerate(targets):
        entry = first_heading.insert_paragraph_before(style=f"TOC {level}")
        _add_internal_hyperlink(entry, text, bookmark)
        entry.paragraph_format.page_break_before = index == split_index
        next_level = targets[index + 1][1] if index + 1 < len(targets) else None
        entry.paragraph_format.keep_with_next = level == 1 and next_level == 2

    first_heading.paragraph_format.page_break_before = True


def format_docx(path: Path) -> None:
    doc = Document(path)

    body = doc.styles["Normal"]
    _set_east_asia_font(body, "PingFang SC")
    body.font.size = Pt(10.5)
    body.paragraph_format.space_after = Pt(4)
    body.paragraph_format.line_spacing = 1.15

    style_sizes = {
        "Title": 24,
        "Subtitle": 13,
        "Author": 10.5,
        "Date": 10.5,
        "Heading 1": 18,
        "Heading 2": 14,
        "Heading 3": 12,
    }
    for style_name, size in style_sizes.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        _set_east_asia_font(style, "PingFang SC")
        style.font.size = Pt(size)
        if style_name.startswith("Heading"):
            style.paragraph_format.keep_with_next = True

    if "Image Caption" in doc.styles:
        caption = doc.styles["Image Caption"]
        _set_east_asia_font(caption, "PingFang SC")
        caption.font.size = Pt(9)
        caption.font.italic = False
        caption.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(6)

    if "Captioned Figure" in doc.styles:
        figure = doc.styles["Captioned Figure"]
        figure.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure.paragraph_format.space_before = Pt(4)
        figure.paragraph_format.space_after = Pt(0)

    _insert_clickable_toc(doc)

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)
        section.header_distance = Cm(0.7)
        section.footer_distance = Cm(0.7)

        header = section.header.paragraphs[0]
        _clear_paragraph(header)
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header.add_run("GIS Data Agent | Gemma 4 决赛技术答辩手册")
        header_run.font.name = "PingFang SC"
        header_run.font.size = Pt(8)
        header_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        footer = section.footer.paragraphs[0]
        _clear_paragraph(footer)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("第 ")
        footer_run.font.name = "PingFang SC"
        footer_run.font.size = Pt(8)
        _add_page_number(footer)
        tail = footer.add_run(" 页")
        tail.font.name = "PingFang SC"
        tail.font.size = Pt(8)

    for table in doc.tables:
        table.autofit = True
        header_text = table.cell(0, 0).text.strip() if table.rows else ""
        compact_scoring_table = header_text == "评审维度"
        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(
                        0 if compact_scoring_table else 1
                    )
                    if compact_scoring_table:
                        paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.name = "PingFang SC"
                        run._element.get_or_add_rPr().get_or_add_rFonts().set(
                            qn("w:eastAsia"), "PingFang SC"
                        )
                        if compact_scoring_table:
                            run.font.size = Pt(7.5 if row_index else 8)
                        else:
                            run.font.size = Pt(8.5 if row_index else 9)

    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    format_docx(args.docx.resolve())


if __name__ == "__main__":
    main()
