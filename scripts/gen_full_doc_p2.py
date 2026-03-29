#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""完整设计文档生成器 - Part 2: 组件架构和数据架构"""
from docx import Document

doc = Document('D:\\adk\\设计文档_完整版_Part1.docx')

# 3. 总体组件/服务架构
doc.add_heading('3. 总体组件/服务架构', 1)
doc.add_paragraph("""核心服务组件清单：""")

# 组件表格
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '组件名称'
hdr[1].text = '端口'
hdr[2].text = '技术栈'
hdr[3].text = '职责'

components = [
    ['Chainlit Server', '8000', 'Chainlit 2.9.6 + Starlette', 'Web UI + WebSocket + 会话管理'],
    ['Intent Router', '-', 'Gemini 2.0 Flash', '语义意图分类 + 多语言检测'],
    ['Pipeline Orchestrator', '-', 'ADK v1.27', '三条管线调度 + Agent编排'],
    ['MCP Hub', '-', 'MCP v1.24', 'stdio/SSE/HTTP协议集成'],
    ['Workflow Engine', '-', 'Python + Cron', 'DAG执行 + 调度 + Webhook'],
    ['PostgreSQL', '5432', 'PostgreSQL 16 + PostGIS 3.4', '主数据库 + 空间数据'],
    ['Redis', '6379', 'Redis 7', '流数据缓存（可选）'],
    ['CV Detection', '8010', 'FastAPI + YOLO', '视觉检测服务'],
    ['CAD Parser', '8011', 'FastAPI + ezdxf', 'CAD/3D解析'],
    ['Reference Data', '8012', 'FastAPI + PostGIS', '参考数据服务'],
]

for comp in components:
    row = table.add_row().cells
    for i, val in enumerate(comp):
        row[i].text = val

doc.add_paragraph("""
服务间通信：
• 同步调用：REST API (HTTP/HTTPS)
• 异步消息：WebSocket (Chainlit)
• 数据库连接：SQLAlchemy连接池（pool_size=5, max_overflow=10）
• MCP通信：stdio子进程、SSE流、HTTP请求""")

# 4. 总体数据架构
doc.add_heading('4. 总体数据架构', 1)
doc.add_paragraph("""采用"湖仓一体"架构，支持结构化、半结构化和非结构化数据统一管理。""")

doc.add_heading('4.1 数据分层', 2)
doc.add_paragraph("""
【原始数据层】
• 用户沙箱：uploads/{user_id}/ 目录隔离
• 支持格式：Shapefile、GeoJSON、GPKG、KML、KMZ、CSV、Excel、TIFF、PDF、DOCX
• 自动检测：坐标系识别、格式转换、ZIP解压

【数据湖层】
• 统一目录：agent_data_assets表
• 四层元数据：
  - Technical：存储位置、空间范围、CRS、要素数量
  - Business：语义标签、分类分级、关键词
  - Operational：创建者、创建时间、版本号
  - Lineage：上游依赖、下游使用、处理历史
• 版本管理：agent_asset_versions表，快照存储
• 血缘追踪：上下游依赖关系图

【语义层】
• 三级架构：
  1. YAML静态目录 (semantic_catalog.yaml)：领域定义、列域、区域分组
  2. 数据库注册表 (agent_semantic_registry)：表/列级语义标注
  3. 自定义域 (agent_semantic_domains)：用户定义层次结构
• 缓存策略：5分钟TTL，写入时失效
• 模糊匹配：同义词、单位后缀、嵌入向量

【应用层】
• PostGIS空间库：GEOMETRY(Point/Polygon, 4326)
• 时序数据：stream_locations表（TimescaleDB-ready）
• 向量索引：pgvector ivfflat，64维L2归一化嵌入
• 知识图谱：agent_kb_entities + agent_kb_relations""")

doc.add_heading('4.2 数据流转', 2)
doc.add_paragraph("""
用户上传 → 格式检测 → 坐标系识别 → 数据剖析 → 语义标注 → 入湖登记 → 管线处理 → 结果输出 → 版本归档

关键节点：
• 格式检测：基于文件扩展名和内容特征
• 坐标系识别：检测常见列名（lng/lat、x/y、longitude/latitude）
• 数据剖析：统计摘要、空间范围、质量评分
• 语义标注：自动匹配 + 用户确认
• 管线处理：根据意图路由到对应管线
• 版本归档：快照存储 + 增量追踪""")

doc.save('D:\\adk\\设计文档_完整版_Part2.docx')
print("Part 2 saved")
