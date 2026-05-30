"""Inspect docx structure: count paragraphs/tables, dump first N items."""
import sys
from pathlib import Path
import docx

def inspect(path: str, head_paragraphs: int = 30, head_tables: int = 3):
    print("="*80)
    print(f"FILE: {path}")
    print("="*80)
    try:
        d = docx.Document(path)
    except Exception as e:
        print(f"[ERROR] cannot open: {e}")
        return

    paras = d.paragraphs
    tables = d.tables
    print(f"paragraphs: {len(paras)}")
    print(f"tables:     {len(tables)}")
    non_empty_paras = [p for p in paras if p.text.strip()]
    print(f"non-empty paragraphs: {len(non_empty_paras)}")

    print()
    print("--- HEADING-LIKE PARAGRAPHS (style starts with Heading) ---")
    headings = [(i, p) for i, p in enumerate(paras) if p.style and p.style.name and p.style.name.startswith("Heading")]
    for i, p in headings[:25]:
        print(f"  [{i:4d}] [{p.style.name}] {p.text.strip()[:80]}")
    print(f"(total headings: {len(headings)})")

    print()
    print(f"--- FIRST {head_paragraphs} NON-EMPTY PARAGRAPHS ---")
    for i, p in enumerate(non_empty_paras[:head_paragraphs]):
        style = p.style.name if p.style else "?"
        print(f"  [{i:3d}] [{style[:15]:15s}] {p.text.strip()[:110]}")

    print()
    print(f"--- FIRST {head_tables} TABLES (shape + first few rows) ---")
    for ti, t in enumerate(tables[:head_tables]):
        rows = t.rows
        if not rows:
            continue
        cols = len(rows[0].cells)
        print(f"  Table #{ti}: {len(rows)} rows × {cols} cols")
        # header
        for ri, row in enumerate(rows[:6]):
            cells = [c.text.strip().replace("\n"," | ")[:30] for c in row.cells]
            print(f"    row {ri}: {cells}")

    if len(tables) > head_tables:
        print()
        print(f"--- TABLE SHAPES (all {len(tables)} tables) ---")
        shapes = []
        for ti, t in enumerate(tables):
            rows = t.rows
            if rows:
                shapes.append((ti, len(rows), len(rows[0].cells)))
        for ti, r, c in shapes:
            print(f"  Table #{ti:3d}: {r:4d} rows × {c} cols")


if __name__ == "__main__":
    inspect(sys.argv[1])
