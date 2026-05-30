"""Find DLTB table in 统一调查监测 docx."""
import docx
from pathlib import Path

p = Path(r"D:\adk\数据标准\自然资源一张图数据库标准1128\自然资源“一张图”数据库体系结构（2）统一调查监测1126.docx")
d = docx.Document(p)

# Find headings/paragraphs mentioning DLTB or 地类图斑
print("--- DLTB / 地类图斑 mentions in paragraphs ---")
for i, para in enumerate(d.paragraphs):
    t = para.text.strip()
    if "地类图斑" in t or "DLTB" in t.upper():
        style = para.style.name if para.style else "?"
        print(f"  [{i:4d}] [{style[:12]:12s}] {t[:120]}")

# Find tables whose attribute name == DLTB or contains "地类图斑"
print()
print("--- Tables containing DLTB cell ---")
for ti, t in enumerate(d.tables):
    rows = t.rows
    if not rows:
        continue
    # Check first 3 rows for DLTB or 地类图斑
    found = False
    for ri in range(min(3, len(rows))):
        for cell in rows[ri].cells:
            txt = cell.text.strip()
            if txt == "DLTB" or "地类图斑" in txt:
                found = True
                break
        if found:
            break
    if found:
        cols = len(rows[0].cells)
        # Print first row + cell that matched
        header = [c.text.strip()[:25] for c in rows[0].cells]
        sample = [c.text.strip()[:25] for c in rows[1].cells] if len(rows) > 1 else []
        print(f"  Table #{ti}: {len(rows)} rows × {cols} cols")
        print(f"    header: {header}")
        print(f"    row1:   {sample}")
