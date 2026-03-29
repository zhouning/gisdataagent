#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""完整设计文档生成器 - 基于实际代码库"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('时空数据中台产品详细设计V3.0.0.0（Data Agent部分）', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. 概述
doc.add_heading('1. 概述', 1)
doc.add_paragraph("""本文档描述时空数据中台产品v3.0.0.0（Data Agent）的系统架构详细设计。

产品定位：基于Google ADK v1.27的企业级地理信息智能分析平台
版本：v15.8 (2026-03)
代码规模：96测试文件、2680+测试用例、202个REST API、48个数据库表、36个工具集
核心能力：数据治理、土地优化（DRL）、空间智能分析、测绘质检
技术特性：多模态融合、自服务扩展、企业可观测性、多租户隔离""")

# 2. 总体技术架构
doc.add_heading('2. 总体技术架构', 1)
doc.add_paragraph("""系统采用分层微服务架构，包含六个核心层次：""")

doc.add_paragraph("""接入层：
• Chainlit UI Server (端口8000)：提供Web界面和WebSocket连接
• REST API Gateway：202个端点，JWT认证
• OAuth2 Provider：支持Google/GitHub第三方登录""")

doc.add_paragraph("""应用层：
• Intent Router：基于Gemini 2.0 Flash的语义意图分类器
• Pipeline Orchestrator：三条专业管线（优化/治理/通用）
• Dynamic Planner Agent：跨管线任务编排
• Workflow Engine：DAG工作流执行器，支持Cron调度""")

doc.add_paragraph("""工具层：
• 36个Toolset模块：空间处理、分析、可视化、治理、融合等
• MCP Hub：支持stdio/SSE/HTTP三种传输协议
• User Tools Engine：用户自定义工具执行引擎""")

doc.add_paragraph("""数据层：
• PostgreSQL 16 + PostGIS 3.4：主数据库（48张表）
• Redis：实时流数据缓存（可选）
• 对象存储：支持Huawei OBS/AWS S3/GCS""")

doc.add_paragraph("""AI服务层：
• Model Gateway：任务感知的模型路由（Gemini 2.0/2.5 Flash/Pro）
• Context Manager：可插拔上下文提供器
• Prompt Registry：版本化提示词管理
• Evaluation Framework：场景化评估框架""")

doc.add_paragraph("""监控层：
• Prometheus Exporter：25+指标
• 结构化日志：JSON格式，trace_id关联
• Alert Engine：阈值告警，Webhook推送""")

doc.save('D:\\adk\\设计文档_Part1.docx')
print("Part 1 done")
