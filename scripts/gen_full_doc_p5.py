#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Part 5: 更多核心模块 + 框架选型"""
from docx import Document

doc = Document('D:\\adk\\设计文档_完整版_Part4.docx')

# 6.3 Frontend模块
doc.add_heading('6.3 Frontend模块', 2)
doc.add_heading('6.3.1 简介', 3)
doc.add_paragraph("""三面板React SPA：Chat（左）| Map（中）| Data（右）

技术栈：React 18 + TypeScript + Vite + Leaflet + deck.gl + ReactFlow
核心特性：可调整面板宽度、2D/3D地图切换、26个数据面板标签页""")

doc.add_heading('6.3.2 关键组件', 3)
doc.add_paragraph("""
• App.tsx：主容器，状态管理，面板布局
• ChatPanel.tsx：消息流，文件上传，语音输入
• MapPanel.tsx：2D（Leaflet）/3D（deck.gl+MapLibre）地图
• DataPanel.tsx：26个标签页（文件、目录、历史、工作流等）
• WorkflowEditor.tsx：ReactFlow DAG编辑器""")

# 6.4 数据库设计
doc.add_heading('6.4 数据库设计', 2)
doc.add_paragraph("""48张表，43个迁移文件""")

# 表汇总
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '类别'
hdr[1].text = '表数量'
hdr[2].text = '代表性表'

categories = [
    ['认证用户', '3', 'agent_app_users, agent_user_memories, agent_table_ownership'],
    ['Token使用', '2', 'agent_token_usage, agent_audit_log'],
    ['协作共享', '4', 'agent_share_links, agent_teams, agent_team_members, agent_map_annotations'],
    ['工作流', '2', 'agent_workflows, agent_workflow_runs'],
    ['语义元数据', '4', 'agent_semantic_registry, agent_semantic_sources, agent_semantic_domains'],
    ['数据资产', '3', 'agent_data_assets, agent_asset_versions, agent_data_requests'],
    ['自定义扩展', '3', 'agent_custom_skills, agent_user_tools, agent_skill_bundles'],
    ['MCP集成', '2', 'agent_mcp_servers, agent_mcp_tool_rules'],
    ['知识库', '4', 'agent_kb_entities, agent_kb_relations, agent_kb_documents'],
    ['质量治理', '3', 'agent_quality_rules, agent_quality_trends, agent_qc_reviews'],
    ['监控告警', '2', 'agent_alert_rules, agent_alert_history'],
    ['其他', '16', '流数据、融合、评估、提示词等'],
]

for cat in categories:
    row = table.add_row().cells
    for i, val in enumerate(cat):
        row[i].text = val

# 7. 系统开发框架选型
doc.add_heading('7. 系统开发框架选型', 1)

doc.add_heading('7.1 后端开发框架', 2)
doc.add_paragraph("""
• 核心框架：Google ADK v1.27（Agent编排）
• Web框架：Chainlit 2.9.6 + Starlette 0.50.0
• API框架：FastAPI 0.123.10（子系统）
• ORM：SQLAlchemy 2.0.45 + asyncpg 0.31.0
• 空间库：GeoPandas 1.1.2 + Shapely 2.1.2 + Rasterio 1.5.0
• AI库：google-genai 1.55.0 + anthropic 0.81.0 + langgraph 1.0.5
• ML库：PyTorch 2.9.1 + Stable-Baselines3 2.7.1
• 总依赖：329个包""")

doc.add_heading('7.2 前端开发框架', 2)
doc.add_paragraph("""
• 核心框架：React 18.3.1 + TypeScript 5.7.2
• 构建工具：Vite 6.3.5
• 状态管理：Recoil 0.7.7 + @chainlit/react-client 0.3.1
• 地图库：Leaflet 1.9.4 + react-map-gl 8.1.0 + maplibre-gl 5.19.0
• 3D渲染：@deck.gl 9.2.10
• 工作流：@xyflow/react 12.10.1
• 图表：ECharts 6.0.0
• 表格：@tanstack/react-table 8.21.3""")

doc.save('D:\\adk\\设计文档_完整版_Part5.docx')
print("Part 5 saved")
