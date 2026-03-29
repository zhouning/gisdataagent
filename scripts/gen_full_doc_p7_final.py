#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Part 7: 中间件 + 依赖 + 数据库详细设计"""
from docx import Document

doc = Document('D:\\adk\\设计文档_完整版_Part6.docx')

# 10. 中间件支持
doc.add_heading('10. 中间件支持', 1)
doc.add_paragraph("""
• Redis 7：实时流数据缓存（可选，默认内存模式）
• MCP Servers：
  - CV Detection Service (FastAPI + YOLO)
  - CAD Parser Service (FastAPI + ezdxf + trimesh)
  - ArcGIS Pro Tools (ArcPy桥接)
  - QGIS Tools (几何验证)
  - Blender Tools (3D处理)""")

# 11. 其它相关依赖情况
doc.add_heading('11. 其它相关依赖情况', 1)

doc.add_heading('11.1 公共研发中心应用底座相关CBB依赖情况', 2)
doc.add_paragraph("""无CBB依赖。本产品为独立部署的AI Agent平台。""")

doc.add_heading('11.2 商业组件相关许可依赖情况', 2)
doc.add_paragraph("""
• Google Gemini API：按Token计费
• Anthropic Claude API：按Token计费（可选）
• 高德地图API：按调用次数计费（可选）
• 天地图Token：免费申请（可选）
• Huawei OBS：按存储和流量计费（可选）

所有商业组件均为可选，系统可在无外部依赖情况下运行（使用本地模型和存储）。""")

# 12. 数据库设计
doc.add_heading('12. 数据库设计', 1)

doc.add_heading('12.1 逻辑模型设计', 2)

doc.add_heading('12.1.1 表汇总', 3)
doc.add_paragraph("""共48张表，分为12个功能域：""")

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '功能域'
hdr[1].text = '表数量'
hdr[2].text = '核心表'
hdr[3].text = '说明'

domains = [
    ['认证用户', '3', 'agent_app_users', 'PBKDF2密码、RBAC角色'],
    ['Token使用', '2', 'agent_token_usage', 'LLM计费、审计日志'],
    ['协作共享', '4', 'agent_teams', '团队、成员、分享链接'],
    ['模板工作流', '3', 'agent_workflows', 'DAG定义、执行历史'],
    ['语义元数据', '4', 'agent_semantic_registry', '三级语义架构'],
    ['数据资产', '6', 'agent_data_assets', '四层元数据、版本、血缘'],
    ['自定义扩展', '3', 'agent_custom_skills', 'Skills、Tools、Bundles'],
    ['MCP集成', '2', 'agent_mcp_servers', '服务器配置、工具规则'],
    ['知识库', '4', 'agent_kb_entities', 'GraphRAG实体关系'],
    ['质量治理', '3', 'agent_quality_rules', '规则、趋势、QC复核'],
    ['监控告警', '2', 'agent_alert_rules', '阈值规则、告警历史'],
    ['其他', '12', 'stream_locations等', '流数据、融合、评估等'],
]

for domain in domains:
    row = table.add_row().cells
    for i, val in enumerate(domain):
        row[i].text = val

doc.add_heading('12.1.2 总体E-R图', 3)
doc.add_paragraph("""
核心实体关系：

用户 (agent_app_users) 1:N 数据资产 (agent_data_assets)
用户 1:N 自定义Skills (agent_custom_skills)
用户 1:N 工作流 (agent_workflows)
用户 N:M 团队 (agent_teams) 通过 agent_team_members
数据资产 1:N 版本 (agent_asset_versions)
数据资产 1:N 访问请求 (agent_data_requests)
工作流 1:N 执行记录 (agent_workflow_runs)
知识库 1:N 实体 (agent_kb_entities)
实体 N:M 实体 通过 agent_kb_relations
MCP服务器 1:N 工具规则 (agent_mcp_tool_rules)
告警规则 1:N 告警历史 (agent_alert_history)

PostGIS空间关系：
stream_configs.geofence (Polygon) 包含 stream_locations.geom (Point)
""")

doc.add_heading('12.1.3 表清单', 3)
doc.add_paragraph("""详见附录：48张表的完整DDL定义（43个迁移文件）。

关键表说明：
• agent_app_users：用户认证，password_hash使用PBKDF2
• agent_data_assets：统一数据目录，四层JSONB元数据
• agent_semantic_registry：列级语义标注，支持同义词
• agent_workflows：工作流定义，steps/parameters为JSONB
• agent_custom_skills：自定义Agent，toolset_names为TEXT[]
• agent_mcp_servers：MCP配置，支持stdio/SSE/HTTP
• agent_kb_entities：知识图谱实体，confidence评分
• agent_quality_trends：质量趋势，6维度评分
• stream_locations：时序位置数据，GEOMETRY(Point, 4326)

RLS策略：25+张表启用行级安全，基于app.current_user上下文变量。""")

doc.save('D:\\adk\\设计文档_完整版_Final.docx')
print("Final document saved")
