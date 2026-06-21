#!/usr/bin/env python3
"""Convert the TWM strategic proposal Markdown into a formatted DOCX file."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "twm-natural-resource-ministry-strategic-technical-proposal.md"
TARGET = ROOT / "docs" / "twm-natural-resource-ministry-strategic-technical-proposal.docx"


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, *, first_line: bool = False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.18
    if first_line:
        pf.first_line_indent = Cm(0.74)


def set_style_font(style, cn_font="宋体", en_font="Times New Roman", size: float | None = None) -> None:
    style.font.name = en_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    if size is not None:
        style.font.size = Pt(size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, 10, bold)


def add_inline_markdown(paragraph, text: str, *, size: float = 10.5, bold_default: bool = False) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, bold_default)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped.startswith("|") and re.fullmatch(r"\|?[\s:\-|]+\|?", stripped))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i])
        i += 1
    rows: list[list[str]] = []
    for line in table_lines:
        if is_table_separator(line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows, i


def build_docx() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    set_style_font(document.styles["Normal"], size=10.5)
    for style_name, font_size in [
        ("Title", 18),
        ("Heading 1", 15),
        ("Heading 2", 13),
        ("Heading 3", 11.5),
    ]:
        set_style_font(document.styles[style_name], cn_font="黑体", size=font_size)

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if image_match:
            image_path = (SOURCE.parent / image_match.group(2)).resolve()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Cm(16.2))
            set_paragraph_format(paragraph)
            i += 1
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_markdown(paragraph, line[2:], size=18, bold_default=True)
            set_paragraph_format(paragraph)
            i += 1
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline_markdown(paragraph, line[3:], size=15, bold_default=True)
            set_paragraph_format(paragraph)
            i += 1
            continue

        if line.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline_markdown(paragraph, line[4:], size=13, bold_default=True)
            set_paragraph_format(paragraph)
            i += 1
            continue

        if line.startswith("|"):
            rows, next_i = parse_table(lines, i)
            if rows:
                table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(len(table.columns)):
                        text_value = row[c_idx] if c_idx < len(row) else ""
                        cell = table.cell(r_idx, c_idx)
                        set_cell_text(cell, text_value, bold=(r_idx == 0))
                        if r_idx == 0:
                            set_cell_shading(cell, "D9EAF7")
                document.add_paragraph()
            i = next_i
            continue

        if line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Cm(0.6)
            add_inline_markdown(paragraph, line[2:], size=10.5)
            set_paragraph_format(paragraph)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, re.sub(r"^\d+\.\s+", "", line), size=10.5)
            set_paragraph_format(paragraph)
            i += 1
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, line[2:], size=10.5)
            set_paragraph_format(paragraph)
            i += 1
            continue

        paragraph = document.add_paragraph()
        if line.startswith("图 "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_markdown(paragraph, line, size=9.5)
        elif line.startswith("日期：") or line.startswith("面向对象："):
            add_inline_markdown(paragraph, line, size=10.5)
        else:
            add_inline_markdown(paragraph, line, size=10.5)
            set_paragraph_format(paragraph, first_line=True)
            i += 1
            continue
        set_paragraph_format(paragraph)
        i += 1

    document.core_properties.title = "面向自然资源治理的地理空间世界模型（TWM）战略技术说明"
    document.core_properties.subject = "TWM 战略技术说明"
    document.core_properties.keywords = "TWM, 地理空间世界模型, 自然资源治理, 国土空间规划, 用途管制"
    document.save(TARGET)


if __name__ == "__main__":
    build_docx()
