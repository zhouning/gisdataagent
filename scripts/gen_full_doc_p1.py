#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""完整设计文档生成器 - 所有章节"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('时空数据中台产品详细设计V3.0.0.0（Data Agent部分）', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. 概述
doc.add_heading('1. 概述', 1)
doc.add_paragraph("""本文档描述时空数据中台产品v3.0.0.0的系统架构详细设计，用于指导产品开发实现。

产品名称：GIS Data Agent (ADK Edition)
版本：v15.8 (2026-03)
技术基础：Google Agent Developer Kit (ADK) v1.27
代码规模：96测试文件、2680+测试用例、202个REST API端点、48个数据库表、36个工具集

核心能力：
• 数据治理：拓扑验证、质量评估、标准合规、缺陷检测
• 土地优化：基于深度强化学习的用地布局优化
• 空间智能：多模态数据融合、语义分析、知识图谱
• 测绘质检：GB/T 24356标准合规，30类缺陷分类，SLA工作流

技术特性：
• 多模态融合：支持矢量、栅格、表格、点云等10种数据源
• 自服务扩展：自定义Skills、User Tools、工作流编排
• 企业可观测性：Prometheus指标、结构化日志、分布式追踪
• 多租户隔离：用户沙箱、RBAC权限、RLS数据隔离""")

# 2. 总体技术架构
doc.add_heading('2. 总体技术架构', 1)
doc.add_paragraph("""系统采用分层微服务架构，包含六个核心层次：""")

doc.add_paragraph("""【接入层】
• Chainlit UI Server (端口8000)：Web界面 + WebSocket实时通信
• REST API Gateway：202个端点，JWT Cookie认证
• OAuth2 Provider：支持Google/GitHub第三方登录（可选）""")

doc.add_paragraph("""【应用层】
• Intent Router：基于Gemini 2.0 Flash的语义意图分类器，支持中英日三语
• Pipeline Orchestrator：三条专业管线
  - Optimization Pipeline：ParallelAgent → Processing → AnalysisQualityLoop → Viz → Summary
  - Governance Pipeline：Exploration → Processing → ReportLoop
  - General Pipeline：Processing → Viz → SummaryLoop
• Dynamic Planner Agent：跨管线任务编排（可选）
• Custom Skills Engine：用户自定义Agent实例化引擎
• Workflow Engine：DAG工作流执行器，支持Cron调度和Webhook""")

doc.add_paragraph("""【工具层】
• 36个Toolset模块：ExplorationToolset、GeoProcessingToolset、AnalysisToolset、VisualizationToolset、GovernanceToolset、DataCleaningToolset、FusionToolset、KnowledgeGraphToolset等
• MCP Hub：Model Context Protocol集成中心，支持stdio/SSE/HTTP三种传输协议
• User Tools Engine：用户自定义工具执行引擎（http_call/sql_query/file_transform/chain）""")

doc.add_paragraph("""【数据层】
• PostgreSQL 16 + PostGIS 3.4：主数据库（48张表，支持空间数据）
• Redis：实时流数据缓存（可选）
• 对象存储适配器：支持Huawei OBS/AWS S3/Google Cloud Storage
• Data Lake Catalog：统一数据资产目录，四层元数据架构""")

doc.add_paragraph("""【AI服务层】
• Model Gateway：任务感知的模型路由（Gemini 2.0 Flash、2.5 Flash、2.5 Pro）
• Context Manager：可插拔上下文提供器，Token预算管理
• Prompt Registry：版本化提示词管理，环境隔离（dev/staging/prod）
• Evaluation Framework：场景化评估框架，自定义指标""")

doc.add_paragraph("""【监控层】
• Prometheus Exporter：25+指标（LLM/Tool/Pipeline/Cache/CircuitBreaker/HTTP）
• 结构化日志：JSON格式，trace_id关联
• Alert Engine：阈值告警，Webhook/WebSocket推送
• Health Check：/health、/ready、/metrics端点""")

doc.save('D:\\adk\\设计文档_完整版_Part1.docx')
print("Part 1 saved")
