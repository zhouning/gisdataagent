"""
智能交互层 — UC-07 生成治理成果报告

基于 model_advisor 的差距分析结果生成 Word 格式的可验收报告。

报告结构：
  1. 概述（源数据、目标标准、分析时间）
  2. 数据现状摘要（字段数、匹配率、坐标系、记录数）
  3. 字段匹配分析（精确匹配、语义匹配、未匹配）
  4. 差距分析（按严重程度分组：高/中/低/信息）
  5. 调整建议清单（必须做/建议做/可选做）
  6. 附录：完整字段匹配表

输出：Word (.docx) 文件
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from data_agent.intelligence.model_advisor import ModelAdjustmentAdvice

logger = logging.getLogger(__name__)

# 严重程度中文映射
_SEVERITY_LABEL = {
    "high": "高（必填字段缺失）",
    "medium": "中（条件必填缺失）",
    "low": "低（可选字段缺失）",
    "info": "信息（多余字段）",
}

_MATCH_TYPE_LABEL = {
    "exact": "精确匹配",
    "semantic": "语义匹配",
    "unmatched": "未匹配",
}

_GAP_TYPE_LABEL = {
    "missing_in_source": "源数据缺失",
    "extra_in_source": "源数据多余",
    "type_mismatch": "类型不匹配",
    "length_mismatch": "长度不匹配",
}


def _set_cell_shading(cell, color_hex: str):
    """设置表格单元格背景色"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_styled_table(doc: Document, headers: list[str], rows: list[list[str]], header_color: str = "2F5496") -> None:
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, header_color)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)


def generate_report(
    advice: ModelAdjustmentAdvice,
    output_path: str | Path | None = None,
    source_record_count: int = 0,
    source_crs: str = "",
) -> Path:
    """
    生成治理成果报告 Word 文档。

    Args:
        advice: model_advisor 的差距分析结果
        output_path: 输出文件路径（None 则自动生成）
        source_record_count: 源数据记录数
        source_crs: 源数据坐标系

    Returns:
        生成的报告文件路径
    """
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = Path(f"治理分析报告_{advice.target_table}_{timestamp}.docx")
    output_path = Path(output_path)

    doc = Document()

    # ===== 标题 =====
    title = doc.add_heading("数据治理分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 1. 概述 =====
    doc.add_heading("1. 概述", level=1)

    overview_data = [
        ["项目", "内容"],
        ["源数据", advice.source_name],
        ["目标标准表", f"{advice.target_table}（{advice.target_table_label}）"],
        ["源数据字段数", str(advice.source_field_count)],
        ["标准要求字段数", str(advice.target_field_count)],
        ["字段匹配率", f"{advice.match_rate:.0%}"],
        ["分析时间", datetime.now().strftime("%Y-%m-%d %H:%M")],
    ]
    if source_record_count:
        overview_data.insert(3, ["源数据记录数", f"{source_record_count:,}"])
    if source_crs:
        crs_short = source_crs[:80] + "..." if len(source_crs) > 80 else source_crs
        overview_data.insert(3, ["坐标系", crs_short])

    table = doc.add_table(rows=len(overview_data), cols=2)
    table.style = "Table Grid"
    for r_idx, (label, value) in enumerate(overview_data):
        table.rows[r_idx].cells[0].text = label
        table.rows[r_idx].cells[1].text = value
        if r_idx == 0:
            for cell in table.rows[r_idx].cells:
                _set_cell_shading(cell, "2F5496")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ===== 2. 匹配分析 =====
    doc.add_heading("2. 字段匹配分析", level=1)

    exact_count = sum(1 for m in advice.matched_fields if m.match_type == "exact")
    semantic_count = sum(1 for m in advice.matched_fields if m.match_type == "semantic")
    unmatched_count = sum(1 for m in advice.matched_fields if m.match_type == "unmatched")

    doc.add_paragraph(
        f"对源数据的 {advice.source_field_count} 个字段与标准 {advice.target_table} "
        f"的 {advice.target_field_count} 个字段进行匹配分析，结果如下："
    )

    _add_styled_table(
        doc,
        ["匹配类型", "数量", "说明"],
        [
            ["精确匹配", str(exact_count), "字段代码完全一致（不区分大小写）"],
            ["语义匹配", str(semantic_count), "字段代码不同但语义等价（如 shape_area ↔ TBMJ）"],
            ["未匹配", str(unmatched_count), "无法在标准中找到对应字段"],
        ],
    )

    # ===== 3. 差距分析 =====
    doc.add_heading("3. 差距分析", level=1)

    high_gaps = [g for g in advice.gaps if g.severity == "high"]
    medium_gaps = [g for g in advice.gaps if g.severity == "medium"]
    low_gaps = [g for g in advice.gaps if g.severity == "low"]
    info_gaps = [g for g in advice.gaps if g.severity == "info"]

    doc.add_paragraph(
        f"共发现 {len(advice.gaps)} 项差距，其中高优先级 {len(high_gaps)} 项、"
        f"中优先级 {len(medium_gaps)} 项、低优先级 {len(low_gaps)} 项、"
        f"信息 {len(info_gaps)} 项。"
    )

    if high_gaps:
        doc.add_heading("3.1 高优先级差距（必须处理）", level=2)
        _add_styled_table(
            doc,
            ["字段代码", "字段名称", "描述", "建议"],
            [[g.field_code, g.field_name, g.description, g.suggestion[:60]] for g in high_gaps],
            header_color="C00000",
        )

    if medium_gaps:
        doc.add_heading("3.2 中优先级差距（建议处理）", level=2)
        _add_styled_table(
            doc,
            ["字段代码", "字段名称", "描述", "建议"],
            [[g.field_code, g.field_name, g.description, g.suggestion[:60]] for g in medium_gaps],
            header_color="BF8F00",
        )

    if low_gaps:
        doc.add_heading("3.3 低优先级差距（可选处理）", level=2)
        # 过滤掉注释内容（field_code 长度过长的是表尾注释被误解析的）
        real_low = [g for g in low_gaps if len(g.field_code) < 20]
        if real_low:
            _add_styled_table(
                doc,
                ["字段代码", "字段名称", "建议"],
                [[g.field_code, g.field_name, g.suggestion[:60]] for g in real_low],
                header_color="4472C4",
            )

    if info_gaps:
        doc.add_heading("3.4 源数据多余字段（评估处理）", level=2)
        doc.add_paragraph(
            "以下字段存在于源数据中但不在当前标准要求范围内。"
            "可能是旧版标准字段、地方扩展字段或自定义字段，建议逐一评估。"
        )
        _add_styled_table(
            doc,
            ["字段代码", "建议"],
            [[g.field_code, g.suggestion[:80]] for g in info_gaps],
            header_color="808080",
        )

    # ===== 4. 调整建议总结 =====
    doc.add_heading("4. 调整建议总结", level=1)

    doc.add_paragraph(
        f"综合以上分析，源数据 {advice.source_name} 与标准 {advice.target_table}"
        f"（{advice.target_table_label}）的匹配率为 {advice.match_rate:.0%}。"
    )

    actions = []
    if high_gaps:
        actions.append(f"【必须】补充 {len(high_gaps)} 个必填字段：{', '.join(g.field_code for g in high_gaps)}")
    if medium_gaps:
        actions.append(f"【建议】评估并补充 {len(medium_gaps)} 个条件必填字段")
    if info_gaps:
        actions.append(f"【评估】确认 {len(info_gaps)} 个多余字段的处置方式（保留/映射/删除）")

    for action in actions:
        doc.add_paragraph(action, style="List Bullet")

    # ===== 附录 =====
    doc.add_heading("附录：完整字段匹配表", level=1)

    _add_styled_table(
        doc,
        ["源字段", "目标字段", "目标字段名称", "匹配类型", "语义组"],
        [
            [
                m.source_field,
                m.target_field or "—",
                m.target_field_name or "—",
                _MATCH_TYPE_LABEL.get(m.match_type, m.match_type),
                m.group_id or "—",
            ]
            for m in advice.matched_fields
        ],
    )

    # ===== 页脚 =====
    doc.add_paragraph()
    footer = doc.add_paragraph("本报告由 GIS Data Agent 自动生成")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 保存
    doc.save(str(output_path))
    logger.info("报告已生成: %s", output_path)
    return output_path
