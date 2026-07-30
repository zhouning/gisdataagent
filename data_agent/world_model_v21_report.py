"""Evidence-backed visual report for county farmland planning runs."""

from __future__ import annotations

import json
import math
import tempfile
import textwrap
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .report_generator import (
    _set_cell_background,
    _setup_page,
    _setup_styles,
    convert_word_to_pdf,
)

PUBLIC_AGENT_NAME = "县域耕地规划 Agent"
PUBLIC_ENGINE_NAME = "县域耕地空间优化引擎"
REPORT_TITLE = "县域耕地空间规划与硬约束审计报告"

NAVY = "17324D"
BLUE = "2563EB"
GREEN = "15803D"
LIGHT_GREEN = "EAF6EE"
RED = "B91C1C"
LIGHT_RED = "FCEBEC"
GRAY = "64748B"
LIGHT_GRAY = "F1F5F9"
WHITE = "FFFFFF"


def _plan_result(context: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    result = context.get("result") or {}
    if not isinstance(result, dict):
        result = {}
    if result.get("mode") == "pipeline_a_to_d" and isinstance(result.get("plan_result"), dict):
        return result, result["plan_result"]
    return result, result


def _load_json(path: Path | None) -> dict[str, Any]:
    if not path or not path.is_file():
        return {}
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def _run_dir(plan: dict[str, Any], audit: dict[str, Any]) -> Path | None:
    for value in (plan.get("out_dir"), audit.get("out_dir")):
        if value:
            path = Path(str(value)).expanduser()
            if path.is_dir():
                return path
    return None


def _artifact_path(
    run_dir: Path | None,
    plan: dict[str, Any],
    *names: str,
) -> Path | None:
    artifacts = plan.get("artifacts") or {}
    candidates: list[Path] = []
    if run_dir:
        for name in names:
            candidates.append(run_dir / name)
    for key in ("map_layer", "optimized_shp", "summary_json"):
        value = artifacts.get(key)
        if not value:
            continue
        path = Path(str(value)).expanduser()
        candidates.append(path)
        if run_dir:
            candidates.append(run_dir / path.name)
    for candidate in candidates:
        if candidate.is_file() and candidate.name in names:
            return candidate
    return None


def _as_float(value: Any) -> float | None:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    return number if math.isfinite(number) else None


def _fmt(value: Any, digits: int = 2, signed: bool = False) -> str:
    number = _as_float(value)
    if number is None:
        return "-"
    prefix = "+" if signed and number > 0 else ""
    if number.is_integer():
        return f"{prefix}{int(number):,}"
    return f"{prefix}{number:,.{digits}f}"


def _duration(value: Any) -> str:
    number = _as_float(value)
    if number is None or number < 0:
        return "-"
    if number < 0.1:
        return f"{number:.3f} 秒"
    return f"{number:.1f} 秒"


def _dataset_label(value: Any) -> str:
    text = str(value or "-")
    return {
        "bishan": "璧山（bishan）",
        "dongxing": "东兴（dongxing）",
    }.get(text, text)


def _configure_plot_fonts() -> None:
    import matplotlib.pyplot as plt

    try:
        from .utils import _configure_fonts

        _configure_fonts()
    except Exception:
        pass
    plt.rcParams["axes.unicode_minus"] = False


def _metric_values(normalized: dict[str, Any], raw_summary: dict[str, Any]) -> dict[str, Any]:
    raw_results = raw_summary.get("results") or []
    raw_first = raw_results[0] if raw_results and isinstance(raw_results[0], dict) else {}
    config = raw_summary.get("config") or {}
    output = raw_summary.get("shapefile_output") or {}
    return {
        "cultivated_area_change_ha": normalized.get(
            "cultivated_area_change_ha", raw_first.get("cultivated_area_change_ha")
        ),
        "slope_change_pct": normalized.get("slope_change_pct", raw_first.get("slope_change_pct")),
        "cont_change": normalized.get("cont_change", raw_first.get("cont_change")),
        "baimu_area_change_ha": normalized.get(
            "baimu_area_change_ha", raw_first.get("baimu_area_change_ha")
        ),
        "total_reward": normalized.get("total_reward", raw_first.get("total_reward")),
        "steps_run": normalized.get("steps_run", raw_first.get("steps_run")),
        "swaps_completed": normalized.get("swaps_completed", raw_first.get("swaps_completed")),
        "n_blocks": normalized.get("n_blocks", config.get("n_blocks")),
        "n_parcels": normalized.get("n_parcels", config.get("n_parcels")),
        "n_input": output.get("n_input"),
        "n_farm_to_forest": output.get("n_farm_to_forest"),
        "n_forest_to_farm": output.get("n_forest_to_farm"),
        "n_unchanged": output.get("n_unchanged"),
        "horizon": config.get("horizon"),
        "top_k": config.get("top_k"),
    }


def _draw_kpi_dashboard(metrics: dict[str, Any], output_path: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    _configure_plot_fonts()
    fig = plt.figure(figsize=(12, 5.5), facecolor="#F8FAFC")
    fig.text(0.04, 0.93, "规划成效概览", fontsize=19, weight="bold", color="#17324D")
    fig.text(
        0.04,
        0.875,
        "所有数值均来自本次 MPC 汇总文件，正负号表示相对规划前的变化",
        fontsize=10,
        color="#64748B",
    )

    cards = [
        (
            "耕地面积",
            f"{_fmt(metrics.get('cultivated_area_change_ha'), 4, True)} ha",
            "满足面积不减少门槛",
            "#EAF6EE",
            "#15803D",
        ),
        (
            "平均坡度",
            f"{_fmt(metrics.get('slope_change_pct'), 4, True)}%",
            "负值表示坡度降低",
            "#EAF6EE",
            "#15803D",
        ),
        (
            "空间连片度",
            _fmt(metrics.get("cont_change"), 4, True),
            "正值表示连片度提升",
            "#EAF6EE",
            "#15803D",
        ),
        (
            "百亩方面积",
            f"{_fmt(metrics.get('baimu_area_change_ha'), 2, True)} ha",
            "集中连片规模增量",
            "#EAF6EE",
            "#15803D",
        ),
    ]
    card_w = 0.215
    gap = 0.02
    for index, (label, value, note, fill, color) in enumerate(cards):
        x = 0.04 + index * (card_w + gap)
        box = FancyBboxPatch(
            (x, 0.47),
            card_w,
            0.31,
            boxstyle="round,pad=0.012,rounding_size=0.015",
            transform=fig.transFigure,
            linewidth=1,
            edgecolor="#D7E2EA",
            facecolor=fill,
        )
        fig.patches.append(box)
        fig.text(x + 0.018, 0.715, label, fontsize=11, color="#475569")
        fig.text(x + 0.018, 0.615, value, fontsize=21, weight="bold", color=color)
        fig.text(x + 0.018, 0.515, note, fontsize=9, color="#64748B")

    scale = [
        ("空间图斑", _fmt(metrics.get("n_parcels"), 0)),
        ("空间块", _fmt(metrics.get("n_blocks"), 0)),
        ("双向置换", f"{_fmt(metrics.get('swaps_completed'), 0)} 对"),
        ("环境步数", _fmt(metrics.get("steps_run"), 0)),
        ("总奖励", _fmt(metrics.get("total_reward"), 2)),
    ]
    for index, (label, value) in enumerate(scale):
        x = 0.04 + index * 0.188
        fig.text(x, 0.31, value, fontsize=15, weight="bold", color="#17324D")
        fig.text(x, 0.245, label, fontsize=9.5, color="#64748B")
    fig.lines.append(
        plt.Line2D([0.04, 0.96], [0.39, 0.39], transform=fig.transFigure, color="#D7E2EA")
    )
    fig.text(
        0.04,
        0.09,
        "决策方法：学习型状态转移模型集成 + 模型预测控制（MPC）",
        fontsize=10.5,
        color="#334155",
    )
    fig.savefig(output_path, dpi=190, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def _flag_code(value: Any) -> int:
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return -1


def _draw_planning_map(
    layer_path: Path | None,
    metrics: dict[str, Any],
    output_path: Path,
) -> bool:
    if not layer_path:
        return False
    try:
        import geopandas as gpd
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np
        from matplotlib.colors import BoundaryNorm, ListedColormap
        from matplotlib.patches import Patch
        from rasterio.features import rasterize
        from rasterio.transform import from_bounds

        _configure_plot_fonts()
        try:
            gdf = gpd.read_file(layer_path, columns=["CHG_FLAG"])
        except (TypeError, ValueError):
            gdf = gpd.read_file(layer_path)
        if gdf.empty or "geometry" not in gdf:
            return False
        gdf = gdf[gdf.geometry.notna() & ~gdf.geometry.is_empty].copy()
        if gdf.empty:
            return False
        flags = gdf.get("CHG_FLAG")
        if flags is None:
            flags = gdf.index.to_series().map(lambda _: -1)
        else:
            flags = flags.map(_flag_code)

        fig = plt.figure(figsize=(12, 7.2), facecolor="#F8FAFC")
        ax = fig.add_axes([0.05, 0.08, 0.62, 0.78])
        ax.set_facecolor("#F8FAFC")
        min_x, min_y, max_x, max_y = gdf.total_bounds
        width = 1400
        data_ratio = (max_y - min_y) / max(max_x - min_x, 1e-12)
        height = max(900, min(1800, int(width * data_ratio)))
        transform = from_bounds(min_x, min_y, max_x, max_y, width, height)
        category_grid = rasterize(
            (
                (geometry, code + 1 if code in {0, 1, 2} else 1)
                for geometry, code in zip(gdf.geometry, flags, strict=False)
            ),
            out_shape=(height, width),
            transform=transform,
            fill=0,
            all_touched=True,
            dtype=np.uint8,
        )
        color_map = ListedColormap(["#F8FAFC", "#DDE4EA", "#DC2626", "#16A34A"])
        norm = BoundaryNorm([-0.5, 0.5, 1.5, 2.5, 3.5], color_map.N)
        ax.imshow(
            category_grid,
            extent=(min_x, max_x, min_y, max_y),
            origin="upper",
            interpolation="nearest",
            cmap=color_map,
            norm=norm,
            aspect="equal",
        )

        ax.set_axis_off()
        fig.text(
            0.04,
            0.94,
            "县域耕地空间布局优化结果",
            fontsize=18,
            weight="bold",
            color="#17324D",
        )
        fig.text(
            0.04,
            0.89,
            "真实优化图层按变化标记分类渲染",
            fontsize=10,
            color="#64748B",
        )
        legend = [
            Patch(facecolor="#DDE4EA", edgecolor="#C4CED8", label="保持不变"),
            Patch(facecolor="#DC2626", edgecolor="#991B1B", label="耕地 -> 林地"),
            Patch(facecolor="#16A34A", edgecolor="#166534", label="林地 -> 耕地"),
        ]
        fig.legend(
            handles=legend,
            loc="upper left",
            bbox_to_anchor=(0.70, 0.76),
            frameon=True,
            framealpha=0.96,
            facecolor="white",
            edgecolor="#D7E2EA",
            fontsize=10,
        )
        fig.text(
            0.70,
            0.48,
            "变化统计",
            fontsize=12,
            weight="bold",
            color="#17324D",
        )
        fig.text(
            0.70,
            0.40,
            f"{_fmt(metrics.get('n_farm_to_forest'), 0)} 个",
            fontsize=18,
            weight="bold",
            color="#B91C1C",
        )
        fig.text(0.70, 0.355, "耕地 -> 林地", fontsize=9.5, color="#64748B")
        fig.text(
            0.84,
            0.40,
            f"{_fmt(metrics.get('n_forest_to_farm'), 0)} 个",
            fontsize=18,
            weight="bold",
            color="#15803D",
        )
        fig.text(0.84, 0.355, "林地 -> 耕地", fontsize=9.5, color="#64748B")
        fig.text(
            0.70,
            0.24,
            f"{_fmt(metrics.get('n_parcels'), 0)} 个",
            fontsize=18,
            weight="bold",
            color="#17324D",
        )
        fig.text(0.70, 0.195, "规划环境图斑", fontsize=9.5, color="#64748B")
        fig.text(
            0.84,
            0.24,
            f"{_fmt(metrics.get('swaps_completed'), 0)} 对",
            fontsize=18,
            weight="bold",
            color="#17324D",
        )
        fig.text(0.84, 0.195, "完成双向置换", fontsize=9.5, color="#64748B")
        fig.savefig(output_path, dpi=210, facecolor=fig.get_facecolor())
        plt.close(fig)
        return True
    except Exception:
        return False


TRACE_LABELS = {
    "world_model_v21_status": "检查版本与运行状态",
    "paper9_inspect_resources": "检查数据与模型资源",
    "paper9_recall_verified_episodes": "召回已验证经验",
    "world_model_v21_pipeline": "选择并执行快速 MPC 路径",
    "world_model_v21_plan": "执行 MPC 空间规划",
    "paper9_audit_run": "执行硬约束审计",
    "paper9_commit_verified_episode": "写入已验证经验库",
}


def _normalized_trace(context: dict[str, Any]) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    for entry in context.get("tool_trace") or []:
        if isinstance(entry, str):
            result.append({"tool_name": entry})
        elif isinstance(entry, dict) and (entry.get("tool_name") or entry.get("name")):
            result.append(
                {
                    "tool_name": str(entry.get("tool_name") or entry.get("name")),
                    "duration_s": entry.get("duration_s", entry.get("duration")),
                }
            )
    return result


def _draw_agent_trace(trace: list[dict[str, Any]], output_path: Path) -> bool:
    if not trace:
        return False
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    _configure_plot_fonts()
    count = len(trace)
    fig, ax = plt.subplots(figsize=(12, 4.6), facecolor="#F8FAFC")
    ax.set_xlim(-0.5, count - 0.5)
    ax.set_ylim(-1.05, 1.0)
    ax.plot(range(count), [0] * count, color="#9FB2C3", linewidth=3, zorder=1)
    colors = ["#2563EB", "#0F766E", "#7C3AED", "#D97706", "#DC2626", "#15803D"]
    for index, entry in enumerate(trace):
        color = colors[index % len(colors)]
        ax.scatter(index, 0, s=420, color=color, edgecolor="white", linewidth=2, zorder=2)
        ax.text(
            index,
            0,
            str(index + 1),
            ha="center",
            va="center",
            color="white",
            weight="bold",
            zorder=3,
        )
        direction = 1 if index % 2 == 0 else -1
        name = str(entry.get("tool_name") or "-")
        wrapped_name = textwrap.fill(name, width=22, break_long_words=True)
        action = TRACE_LABELS.get(name, "执行领域函数")
        y_action = 0.42 * direction
        y_name = 0.62 * direction
        va = "bottom" if direction > 0 else "top"
        ax.text(
            index,
            y_action,
            action,
            ha="center",
            va=va,
            fontsize=9.4,
            weight="bold",
            color="#17324D",
        )
        ax.text(
            index,
            y_name,
            wrapped_name,
            ha="center",
            va=va,
            fontsize=7.2,
            color="#64748B",
        )
        duration_y = -0.12 if direction > 0 else 0.12
        ax.text(
            index,
            duration_y,
            _duration(entry.get("duration_s")),
            ha="center",
            va="center",
            fontsize=8,
            color="#334155",
        )
    ax.set_title(
        "Gemma 4 + Google ADK 原生函数调用轨迹",
        loc="left",
        fontsize=17,
        weight="bold",
        color="#17324D",
    )
    ax.text(
        0,
        0.91,
        "模型负责观察、选择下一工具与受控闭环；确定性代码负责规划计算和硬约束裁决",
        transform=ax.transAxes,
        fontsize=9.5,
        color="#64748B",
    )
    ax.set_axis_off()
    fig.tight_layout()
    fig.savefig(output_path, dpi=190, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return True


def _set_cell_text(
    cell,
    text: Any,
    *,
    bold: bool = False,
    color: str = "334155",
    size: float = 9.5,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")


def _set_table_borders(table, color: str = "D7E2EA", size: str = "4") -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def _table(doc: Document, headers: list[str], rows: list[list[Any]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_background(cell, NAVY)
        _set_cell_text(cell, header, bold=True, color=WHITE, size=9)
        if widths:
            cell.width = Cm(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                _set_cell_background(cells[index], "F8FAFC")
            _set_cell_text(cells[index], value, size=8.8)
            if widths:
                cells[index].width = Cm(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)


def _paragraph(doc: Document, text: str, *, color: str = "334155", size: float = 9.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def _add_cover(
    doc: Document,
    *,
    dataset: str,
    audit_passed: bool,
    committed: bool,
    author: str,
    episode_id: str,
) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_background(cell, NAVY)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.margin_top = Cm(1.2)
    cell.margin_bottom = Cm(1.2)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(REPORT_TITLE)
    run.bold = True
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor.from_string(WHITE)
    p = cell.add_paragraph(PUBLIC_AGENT_NAME)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor.from_string("BFD7EA")
    p.paragraph_format.space_after = Pt(36)

    status_ok = audit_passed and committed
    status_text = (
        "审计通过，已写入已验证经验库"
        if status_ok
        else ("审计通过，尚未写入经验库" if audit_passed else "审计未通过，结果未提交")
    )
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    status_cell = banner.cell(0, 0)
    _set_cell_background(status_cell, LIGHT_GREEN if status_ok else LIGHT_RED)
    _set_cell_text(
        status_cell,
        status_text,
        bold=True,
        color=GREEN if status_ok else RED,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    meta_rows = [
        ["规划范围", dataset, "报告生成", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["生成主体", author or "GIS Data Agent", "经验编号", episode_id or "-"],
    ]
    _table(doc, ["项目", "内容", "项目", "内容"], meta_rows, [2.4, 4.7, 2.4, 4.7])
    _paragraph(
        doc,
        "本报告由运行产物自动生成。图件、指标、函数耗时和审计结论均绑定本次执行记录，可用于方案复核与演示留档。",
        color=GRAY,
        size=9,
    )
    doc.add_page_break()


def _add_picture(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(7)
    caption_run = caption_p.add_run(caption)
    caption_run.font.size = Pt(8.5)
    caption_run.font.color.rgb = RGBColor.from_string(GRAY)


def _audit_rows(audit: dict[str, Any], metrics: dict[str, Any]) -> list[list[str]]:
    policy = audit.get("policy") or {}
    cultivated = _as_float(metrics.get("cultivated_area_change_ha"))
    cultivated_floor = _as_float(policy.get("cultivated_area_floor_delta_ha"))
    slope = _as_float(metrics.get("slope_change_pct"))
    contiguity = _as_float(metrics.get("cont_change"))
    outputs = bool(audit.get("all_expected_outputs_exist"))

    def verdict(condition: bool) -> str:
        return "通过" if condition else "未通过"

    return [
        [
            "耕地面积不减少",
            f">= {_fmt(cultivated_floor, 4)} ha",
            f"{_fmt(cultivated, 4, True)} ha",
            verdict(
                cultivated is not None
                and cultivated_floor is not None
                and cultivated >= cultivated_floor
            ),
        ],
        [
            "平均坡度改善",
            "变化值 < 0",
            f"{_fmt(slope, 4, True)}%",
            verdict(slope is not None and slope < 0),
        ],
        [
            "空间连片度改善",
            "变化值 > 0",
            _fmt(contiguity, 4, True),
            verdict(contiguity is not None and contiguity > 0),
        ],
        ["成果文件完整", "摘要 + 空间结果", "完整" if outputs else "缺失", verdict(outputs)],
    ]


def _build_document(
    context: dict[str, Any],
    *,
    author: str,
    asset_dir: Path,
) -> Document:
    result, plan = _plan_result(context)
    audit = context.get("audit_result") or {}
    commit = context.get("commit_result") or {}
    status = context.get("status_result") or {}
    args = context.get("tool_args") or {}
    trace = _normalized_trace(context)
    run_dir = _run_dir(plan, audit)
    summary_path = _artifact_path(run_dir, plan, "mpc_summary.json")
    raw_summary = _load_json(summary_path)
    normalized = plan.get("summary") or {}
    metrics = _metric_values(normalized, raw_summary)
    layer_path = _artifact_path(
        run_dir,
        plan,
        "optimized_dltb.fgb",
        "optimized_dltb.shp",
        "DLTB_optimized.shp",
    )

    audit_passed = bool(audit.get("hard_constraint_passed"))
    committed = commit.get("status") == "committed"
    episode = commit.get("episode") or {}
    episode_id = str(episode.get("episode_id") or "-")
    dataset = _dataset_label(args.get("dataset") or context.get("dataset"))

    kpi_path = asset_dir / "planning_kpis.png"
    map_path = asset_dir / "planning_map.png"
    trace_path = asset_dir / "agent_trace.png"
    _draw_kpi_dashboard(metrics, kpi_path)
    map_ready = _draw_planning_map(layer_path, metrics, map_path)
    trace_ready = _draw_agent_trace(trace, trace_path)

    doc = Document()
    _setup_styles(doc)
    _setup_page(doc)
    doc.core_properties.title = REPORT_TITLE
    doc.core_properties.author = author or "GIS Data Agent"
    doc.core_properties.subject = "县域耕地空间规划、硬约束审计与已验证经验"

    _add_cover(
        doc,
        dataset=dataset,
        audit_passed=audit_passed,
        committed=committed,
        author=author,
        episode_id=episode_id,
    )

    _heading(doc, "1. 执行摘要")
    outcome = (
        "规划方案通过硬约束校验，并在审计通过后写入已验证经验库。"
        if audit_passed and committed
        else "规划结果未完成审计与经验提交闭环，不应作为已验证方案使用。"
    )
    _paragraph(doc, outcome, color=GREEN if audit_passed and committed else RED, size=10.5)
    _add_picture(doc, kpi_path, 14.5, "图 1  本次运行关键规划指标")
    planning_seconds = sum(
        _as_float(entry.get("duration_s")) or 0.0
        for entry in trace
        if entry.get("tool_name") in {"world_model_v21_pipeline", "world_model_v21_plan"}
    )
    governance_seconds = sum(
        _as_float(entry.get("duration_s")) or 0.0
        for entry in trace
        if entry.get("tool_name") in {"paper9_audit_run", "paper9_commit_verified_episode"}
    )
    overview_rows = [
        [
            "数据集",
            dataset,
            "环境",
            "县域" if plan.get("env_kind") == "county" else str(plan.get("env_kind") or "-"),
        ],
        ["规划引擎", PUBLIC_ENGINE_NAME, "算法", "学习型状态转移模型集成 + MPC"],
        [
            "MPC 前瞻步长",
            _fmt(args.get("horizon") or metrics.get("horizon"), 0),
            "每步候选行动数",
            _fmt(args.get("top_k") or metrics.get("top_k"), 0),
        ],
        [
            "总用时",
            _duration(context.get("total_duration_s")),
            "MPC 规划用时",
            _duration(planning_seconds),
        ],
        ["原生函数调用", f"{len(trace)} 次", "审计与经验提交", _duration(governance_seconds)],
    ]
    _table(doc, ["项目", "数值", "项目", "数值"], overview_rows, [2.5, 4.5, 2.5, 4.5])
    doc.add_page_break()

    _heading(doc, "2. 空间规划结果")
    _paragraph(
        doc,
        "图件直接读取本次运行的优化后空间图层。灰色为保持不变，红色为耕地转林地，绿色为林地转耕地。双向置换在县域尺度上保持耕地总量门槛，同时优化坡度与连片度。",
    )
    if map_ready:
        _add_picture(doc, map_path, 14.5, "图 2  县域耕地空间布局优化变化图")
    else:
        warning = doc.add_table(rows=1, cols=1).cell(0, 0)
        _set_cell_background(warning, LIGHT_RED)
        _set_cell_text(
            warning,
            "空间图层未能渲染，报告保留指标和审计证据，请复核原始交付物。",
            color=RED,
            bold=True,
        )
    transition_rows = [
        [
            "输入记录",
            _fmt(metrics.get("n_input"), 0),
            "进入规划环境的图斑",
            _fmt(metrics.get("n_parcels"), 0),
        ],
        [
            "耕地 -> 林地",
            _fmt(metrics.get("n_farm_to_forest"), 0),
            "林地 -> 耕地",
            _fmt(metrics.get("n_forest_to_farm"), 0),
        ],
        [
            "保持不变",
            _fmt(metrics.get("n_unchanged"), 0),
            "完成双向置换",
            f"{_fmt(metrics.get('swaps_completed'), 0)} 对",
        ],
    ]
    _table(doc, ["统计项", "数量", "统计项", "数量"], transition_rows, [3.1, 3.9, 3.1, 3.9])
    doc.add_page_break()

    _heading(doc, "3. Gemma 4 + Google ADK 受控自主闭环")
    _paragraph(
        doc,
        "Gemma 4 根据版本、资源、历史经验与审计反馈选择下一工具；"
        "Google ADK 执行结构化原生函数调用；确定性规划引擎完成空间搜索；"
        "硬约束校验决定是否允许经验写入。",
    )
    if trace_ready:
        _add_picture(doc, trace_path, 14.5, "图 3  本次运行的原生函数调用顺序与真实耗时")
    trace_rows = []
    for index, entry in enumerate(trace, start=1):
        name = str(entry.get("tool_name") or "-")
        trace_rows.append(
            [
                index,
                TRACE_LABELS.get(name, "执行领域函数"),
                name,
                _duration(entry.get("duration_s")),
            ]
        )
    if trace_rows:
        _table(doc, ["序号", "观察/决策", "原生函数", "用时"], trace_rows, [1.2, 4.4, 6.2, 2.2])

    stage_labels = {
        "prepare": "A / 数据准备",
        "sample": "B / 样本生成",
        "train": "C / 状态转移模型训练",
        "plan": "D / MPC 规划执行",
    }
    stage_status = {"skipped_reused": "已复用", "ok": "完成", "ready": "就绪"}
    stage_rows = []
    for step in result.get("steps") or []:
        if not isinstance(step, dict):
            continue
        key = str(step.get("step") or "-")
        value = str(step.get("status") or step.get("mode") or "-")
        stage_rows.append([stage_labels.get(key, key), stage_status.get(value, value)])
    if stage_rows:
        _heading(doc, "算法执行阶段", level=2)
        _table(doc, ["阶段", "本次状态"], stage_rows, [9.5, 4.5])
    doc.add_page_break()

    _heading(doc, "4. 硬约束审计与经验治理")
    audit_banner = doc.add_table(rows=1, cols=1).cell(0, 0)
    _set_cell_background(audit_banner, LIGHT_GREEN if audit_passed else LIGHT_RED)
    _set_cell_text(
        audit_banner,
        f"硬约束审计：{'通过' if audit_passed else '未通过'}",
        color=GREEN if audit_passed else RED,
        bold=True,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(5)
    _table(
        doc,
        ["审计项", "约束", "实际结果", "结论"],
        _audit_rows(audit, metrics),
        [4.3, 3.3, 3.8, 2.6],
    )

    memory_text = "已写入" if committed else "未写入"
    memory_rows = [
        ["已验证经验库", memory_text],
        ["经验编号", episode_id],
        ["写入规则", "仅当硬约束审计通过时允许写入"],
    ]
    _heading(doc, "经验写入", level=2)
    _table(doc, ["项目", "结果"], memory_rows, [4.1, 9.9])

    engine = status.get("paper9") or {}
    summary_artifact = (audit.get("artifacts") or {}).get("summary") or {}
    checksum = str(summary_artifact.get("sha256") or "-")
    deliverable_rows = [
        [
            "规划汇总",
            "JSON",
            "关键指标与运行参数",
            checksum[:16] + ("..." if len(checksum) > 16 else ""),
        ],
        [
            "优化空间结果",
            "FlatGeobuf / Shapefile",
            "地图复核与 GIS 交付",
            "已生成" if layer_path else "缺失",
        ],
        ["硬约束审计记录", "JSON", "校验结论与失败原因", "已生成" if audit else "缺失"],
        ["已验证经验", "Append-only 记录", "后续任务受控召回", memory_text],
    ]
    _heading(doc, "可追溯交付", level=2)
    _table(doc, ["交付物", "格式", "用途", "状态/摘要哈希"], deliverable_rows, [3.4, 3.2, 4.2, 3.2])
    _paragraph(
        doc,
        (
            "版本信息：适配器 "
            f"{plan.get('version', result.get('version', '-'))}；"
            f"引擎包 {engine.get('package_version', '-')}；"
            f"算法 {engine.get('algorithm_version', '-')}。"
        ),
        color=GRAY,
        size=8.5,
    )
    _paragraph(
        doc,
        "边界说明：本报告记录计算与工程审计结果，不替代法定规划审批、权属确认或生产数据验收。",
        color=GRAY,
        size=8.5,
    )
    return doc


def generate_world_model_v21_word_report(
    context: dict[str, Any],
    output_path: str,
    *,
    author: str = "",
) -> str:
    """Generate a visual DOCX from one structured county-planning run."""
    output = Path(output_path).expanduser().resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="wm_v21_report_") as tmp:
        doc = _build_document(context, author=author, asset_dir=Path(tmp))
        doc.save(output)
    return str(output)


def generate_world_model_v21_pdf_report(
    context: dict[str, Any],
    output_path: str,
    *,
    author: str = "",
) -> str:
    """Generate a visual PDF from one structured county-planning run."""
    output = Path(output_path).expanduser().resolve()
    docx_path = output.with_suffix(".docx")
    generate_world_model_v21_word_report(context, str(docx_path), author=author)
    return convert_word_to_pdf(str(docx_path), str(output), remove_source=True)
