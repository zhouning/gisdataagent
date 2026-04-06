"""
知识管理层 — 标准规则库

从自然资源"一张图"数据库标准 Word 文档中解析出结构化规则条目。
文档格式：Word (.docx)，内含两类表格：
  1. 图层定义表：序号|层名|层要素|几何特征|属性表名|约束条件|说明
  2. 字段定义表：序号|字段名称|字段代码|字段类型|字段长度|小数位数|值域|约束条件|备注

Phase 1: 解析三调对应的标准文档（统一调查监测）
"""

from __future__ import annotations

import logging
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class FieldRule:
    """数据标准中的字段规则"""

    seq: int  # 序号
    field_name: str  # 字段名称（中文）
    field_code: str  # 字段代码（英文缩写，如 DLBM）
    field_type: str  # 字段类型（Char/Float/Int/Date/...）
    field_length: int | None = None  # 字段长度
    decimal_places: int | None = None  # 小数位数
    value_domain: str = ""  # 值域（如"见表5-35"或"[0,1)"）
    constraint: str = ""  # 约束条件（M=必填, O=可选, C=条件必填）
    remark: str = ""  # 备注


@dataclass
class LayerRule:
    """数据标准中的图层规则"""

    seq: int
    layer_group: str  # 层名
    feature_name: str  # 层要素
    geometry_type: str  # 几何特征（Point/Polygon/Line/Annotation）
    table_name: str  # 属性表名（如 DLTB）
    constraint: str = ""  # 约束条件
    remark: str = ""


@dataclass
class TableStandard:
    """一个属性表的完整标准定义"""

    table_name: str  # 属性表名（如 DLTB）
    table_label: str = ""  # 中文名（如"地类图斑"）
    fields: list[FieldRule] = field(default_factory=list)

    def get_field(self, code: str) -> FieldRule | None:
        """按字段代码查找"""
        for f in self.fields:
            if f.field_code == code:
                return f
        return None

    def mandatory_fields(self) -> list[FieldRule]:
        """返回所有必填字段"""
        return [f for f in self.fields if f.constraint == "M"]

    def to_dict(self) -> dict:
        """导出为字典结构（供 AI 推理使用）"""
        return {
            "属性表名": self.table_name,
            "中文名": self.table_label,
            "字段": [
                {
                    "序号": f.seq,
                    "字段名称": f.field_name,
                    "字段代码": f.field_code,
                    "字段类型": f.field_type,
                    "字段长度": f.field_length,
                    "小数位数": f.decimal_places,
                    "值域": f.value_domain,
                    "约束条件": f.constraint,
                    "备注": f.remark,
                }
                for f in self.fields
            ],
        }


@dataclass
class StandardDocument:
    """一份标准文档的完整解析结果"""

    standard_name: str
    layers: list[LayerRule] = field(default_factory=list)
    tables: list[TableStandard] = field(default_factory=list)

    def get_table(self, table_name: str) -> TableStandard | None:
        """按属性表名查找"""
        for t in self.tables:
            if t.table_name == table_name:
                return t
        return None

    def summary(self) -> dict:
        return {
            "标准名称": self.standard_name,
            "图层数": len(self.layers),
            "属性表数": len(self.tables),
            "属性表列表": [
                {
                    "表名": t.table_name,
                    "中文名": t.table_label,
                    "字段数": len(t.fields),
                    "必填字段数": len(t.mandatory_fields()),
                }
                for t in self.tables
            ],
        }

    def to_dict(self) -> dict:
        return {
            "标准名称": self.standard_name,
            "图层": [
                {
                    "层名": l.layer_group,
                    "层要素": l.feature_name,
                    "几何特征": l.geometry_type,
                    "属性表名": l.table_name,
                    "约束条件": l.constraint,
                }
                for l in self.layers
            ],
            "属性表": [t.to_dict() for t in self.tables],
        }


def _safe_int(value: str) -> int | None:
    """安全转换为整数"""
    value = value.strip()
    if not value:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def _is_layer_table(header_cells: list[str]) -> bool:
    """判断是否为图层定义表"""
    return "层名" in header_cells and "属性表名" in header_cells


def _is_field_table(header_cells: list[str]) -> bool:
    """判断是否为字段定义表"""
    return "字段名称" in header_cells and "字段代码" in header_cells


def _find_table_label(doc, table_index: int) -> str:
    """从表格前面的段落中提取表的中文名称。

    标准文档中，表格前通常有一个标题段落，如：
    "表5-13地类图斑属性结构描述表(属性表名：DLTB)"
    """
    # python-docx 没有直接的 table-to-paragraph 映射
    # 用简单启发式：在 doc.element 中查找 table 之前的段落
    from docx.oxml.ns import qn

    table_elem = doc.tables[table_index]._tbl
    prev = table_elem.getprevious()
    while prev is not None:
        if prev.tag == qn("w:p"):
            text = prev.text or ""
            # 尝试从 runs 中提取文本
            if not text:
                runs = prev.findall(f".//{qn('w:t')}")
                text = "".join(r.text or "" for r in runs)
            text = text.strip()
            if "属性表名" in text or "属性结构" in text:
                # 提取中文名：找"XXX属性结构"中的 XXX
                for keyword in ["属性结构描述表", "属性结构"]:
                    if keyword in text:
                        idx = text.index(keyword)
                        # 向前找表名
                        prefix = text[:idx]
                        # 去掉表编号（如"表5-13"）
                        for sep in ["表", "："]:
                            if sep in prefix:
                                prefix = prefix[prefix.rindex(sep) + 1 :]
                        return prefix.strip()
            if text:
                break
        prev = prev.getprevious()
    return ""


def parse_standard_docx(doc) -> StandardDocument:
    """
    解析标准 Word 文档，提取图层定义和字段定义。

    Args:
        doc: python-docx Document 对象

    Returns:
        StandardDocument
    """
    # 提取标准名称（通常在文档开头的段落中）
    standard_name = ""
    for para in doc.paragraphs[:20]:
        text = para.text.strip()
        if "数据库体系结构" in text or "一张图" in text:
            standard_name = text
            break
    if not standard_name:
        standard_name = "未知标准"

    layers: list[LayerRule] = []
    tables: list[TableStandard] = []

    for t_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue

        header_cells = [cell.text.strip() for cell in table.rows[0].cells]

        if _is_layer_table(header_cells):
            # 解析图层定义表
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 5 or not cells[0]:
                    continue
                layers.append(
                    LayerRule(
                        seq=_safe_int(cells[0]) or 0,
                        layer_group=cells[1],
                        feature_name=cells[2],
                        geometry_type=cells[3],
                        table_name=cells[4],
                        constraint=cells[5] if len(cells) > 5 else "",
                        remark=cells[6] if len(cells) > 6 else "",
                    )
                )

        elif _is_field_table(header_cells):
            # 解析字段定义表
            fields: list[FieldRule] = []
            table_name = ""

            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 5:
                    continue
                # 跳过空行或合并行
                if not cells[1] and not cells[2]:
                    continue

                fields.append(
                    FieldRule(
                        seq=_safe_int(cells[0]) or 0,
                        field_name=cells[1],
                        field_code=cells[2],
                        field_type=cells[3],
                        field_length=_safe_int(cells[4]),
                        decimal_places=_safe_int(cells[5])
                        if len(cells) > 5
                        else None,
                        value_domain=cells[6] if len(cells) > 6 else "",
                        constraint=cells[7] if len(cells) > 7 else "",
                        remark=cells[8] if len(cells) > 8 else "",
                    )
                )

            if fields:
                # 尝试从前面的段落提取表名
                label = _find_table_label(doc, t_idx)

                # 尝试从前面的 Caption 段落提取属性表名
                table_elem = table._tbl
                prev = table_elem.getprevious()
                from docx.oxml.ns import qn

                while prev is not None:
                    if prev.tag == qn("w:p"):
                        runs = prev.findall(f".//{qn('w:t')}")
                        text = "".join(r.text or "" for r in runs).strip()
                        if "属性表名" in text:
                            # 提取 "属性表名：XXXX" 中的 XXXX
                            idx = text.index("属性表名")
                            after = text[idx + 4 :].strip().strip("：:（()）)")
                            # 取第一个单词（英文表名）
                            table_name = after.split("/")[0].split("）")[0].split(")")[0].strip()
                            break
                        if text:
                            break
                    prev = prev.getprevious()

                if not table_name:
                    table_name = f"TABLE_{t_idx}"

                tables.append(
                    TableStandard(
                        table_name=table_name,
                        table_label=label,
                        fields=fields,
                    )
                )

    logger.info(
        "标准解析完成: %s, 图层=%d, 属性表=%d",
        standard_name,
        len(layers),
        len(tables),
    )

    return StandardDocument(
        standard_name=standard_name,
        layers=layers,
        tables=tables,
    )


def load_from_zip(
    zip_path: str | Path, docx_filename: str | None = None
) -> StandardDocument:
    """
    从 ZIP 包中加载标准 Word 文档并解析。

    Args:
        zip_path: ZIP 文件路径
        docx_filename: ZIP 内的 docx 文件名（None 则自动匹配"统一调查监测"）
    """
    from docx import Document

    zip_path = Path(zip_path)
    if not zip_path.exists():
        raise FileNotFoundError(f"ZIP 文件不存在: {zip_path}")

    with zipfile.ZipFile(zip_path, "r") as z:
        if docx_filename is None:
            matches = [n for n in z.namelist() if "统一调查监测" in n]
            if not matches:
                raise FileNotFoundError(
                    f"ZIP 中未找到统一调查监测文档，可用文件: {z.namelist()}"
                )
            docx_filename = matches[0]

        with z.open(docx_filename) as f:
            doc = Document(f)

    return parse_standard_docx(doc)


def load_survey_standard(
    zip_path: str | Path = r"D:\自然资源一张图数据库标准1128.zip",
) -> StandardDocument:
    """
    加载三调（统一调查监测）数据标准的便捷方法。
    """
    return load_from_zip(zip_path)
