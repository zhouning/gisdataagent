"""Checks for the generated Abu Dhabi customer delivery documents."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = ROOT / "docs/customer/abu_dhabi_liveability_site_validation"
XLSX = OUTPUT_ROOT / (
    "abu_dhabi_stormwater_world_model_authoritative_data_requirements.xlsx"
)
DOCX = OUTPUT_ROOT / (
    "abu_dhabi_stormwater_world_model_authoritative_data_requirements.docx"
)


def test_authoritative_data_workbook_has_register_rows_and_customer_fill_template():
    workbook = load_workbook(XLSX, read_only=False)
    assert workbook.sheetnames == [
        "权威数据需求清单",
        "交付原则与状态",
        "模型与数据对应关系",
        "客户交付填报模板",
    ]
    requirements = workbook["权威数据需求清单"]
    assert requirements.max_row == 16
    assert requirements.max_column == 18
    assert all(requirements.cell(row, 17).value == "待客户提供" for row in range(5, 17))
    assert requirements.freeze_panes == "A5"
    assert requirements.auto_filter.ref == "A4:R16"
    template = workbook["客户交付填报模板"]
    assert template.max_row == 24
    assert template.max_column == 16
    assert template.freeze_panes == "A5"


def test_customer_word_brief_states_authoritative_only_boundary_and_model_roles():
    document = Document(DOCX)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    combined = f"{text}\n{table_text}"
    assert len(document.tables) == 6
    assert "客户或主管部门提供权威真实数据" in combined
    assert "只用于测试数据管线" in combined
    assert "EPA SWMM 5.2.4" in combined
    assert "ANUGA 2D" in combined
    assert "GWM" in combined
    assert "所有 12 项均保持待客户提供" in combined
    assert "K0、传统模型准入、GWM 训练" in combined
