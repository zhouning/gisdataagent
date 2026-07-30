import zipfile
from pathlib import Path

import geopandas as gpd
from docx import Document
from shapely.geometry import box

from data_agent.world_model_v21_report import generate_world_model_v21_word_report


def _report_context(run_dir: Path) -> dict:
    return {
        "result": {
            "mode": "pipeline_a_to_d",
            "version": "2.1.0",
            "steps": [
                {"step": "prepare", "status": "skipped_reused"},
                {"step": "sample", "status": "skipped_reused"},
                {"step": "train", "status": "skipped_reused"},
                {"step": "plan", "status": "ok"},
            ],
            "plan_result": {
                "status": "ok",
                "version": "2.1.0",
                "env_kind": "county",
                "out_dir": str(run_dir),
                "summary": {
                    "steps_run": 100,
                    "swaps_completed": 2,
                    "n_blocks": 3,
                    "n_parcels": 3,
                    "cultivated_area_change_ha": 1.25,
                    "slope_change_pct": -0.61,
                    "cont_change": 0.02,
                    "baimu_area_change_ha": 4.5,
                    "total_reward": 9.8,
                },
                "artifacts": {
                    "summary_json": "mpc_summary.json",
                    "map_layer": "optimized_dltb.shp",
                },
            },
        },
        "tool_args": {"dataset": "bishan", "horizon": "1", "top_k": "1"},
        "status_result": {"paper9": {"package_version": "0.3.3", "algorithm_version": "2.2.3"}},
        "audit_result": {
            "hard_constraint_passed": True,
            "all_expected_outputs_exist": True,
            "out_dir": str(run_dir),
            "policy": {"cultivated_area_floor_delta_ha": 0.0},
            "artifacts": {"summary": {"sha256": "a" * 64}},
        },
        "commit_result": {
            "status": "committed",
            "episode": {"episode_id": "episode-verified-001"},
        },
        "tool_trace": [
            {"tool_name": "world_model_v21_status", "duration_s": 0.01},
            {"tool_name": "paper9_inspect_resources", "duration_s": 0.02},
            {"tool_name": "paper9_recall_verified_episodes", "duration_s": 0.03},
            {"tool_name": "world_model_v21_pipeline", "duration_s": 12.3},
            {"tool_name": "paper9_audit_run", "duration_s": 0.12},
            {"tool_name": "paper9_commit_verified_episode", "duration_s": 0.04},
        ],
        "total_duration_s": 13.1,
    }


def test_world_model_v21_report_embeds_real_map_and_visual_evidence(tmp_path):
    run_dir = tmp_path / "run"
    run_dir.mkdir()
    gdf = gpd.GeoDataFrame(
        {"CHG_FLAG": [0, 1, 2]},
        geometry=[box(0, 0, 1, 1), box(1, 0, 2, 1), box(0, 1, 1, 2)],
        crs="EPSG:4326",
    )
    gdf.to_file(run_dir / "optimized_dltb.shp")
    (run_dir / "mpc_summary.json").write_text(
        """{
          "config": {"horizon": 1, "top_k": 1, "n_blocks": 3, "n_parcels": 3},
          "results": [{
            "steps_run": 100,
            "swaps_completed": 2,
            "cultivated_area_change_ha": 1.25,
            "slope_change_pct": -0.61,
            "cont_change": 0.02,
            "baimu_area_change_ha": 4.5,
            "total_reward": 9.8
          }],
          "shapefile_output": {
            "n_input": 3,
            "n_in_env": 3,
            "n_farm_to_forest": 1,
            "n_forest_to_farm": 1,
            "n_unchanged": 1
          }
        }""",
        encoding="utf-8",
    )

    output = tmp_path / "county_report.docx"
    report_path = generate_world_model_v21_word_report(
        _report_context(run_dir), str(output), author="Finals Tester"
    )

    document = Document(report_path)
    visible_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "县域耕地空间规划与硬约束审计报告" in visible_text
    assert "Gemma 4 + Google ADK 受控自主闭环" in visible_text
    assert "硬约束审计：通过" in visible_text
    assert "已验证经验库" in visible_text
    assert "episode-verified-001" in visible_text
    assert "/tmp/" not in visible_text

    with zipfile.ZipFile(report_path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert len(media) >= 3
