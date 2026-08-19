from __future__ import annotations

import io
from datetime import UTC, datetime
from pathlib import Path
from unittest.mock import MagicMock

from docx import Document

from data_agent.ontology.authority_reader import PostgresOntologyReader
from data_agent.ontology.cli import build_parser
from data_agent.ontology.compiler import (
    CompiledOntology,
    EAInput,
    StandardInput,
    _parse_standard_document,
    compile_ontology,
    write_package,
)
from data_agent.ontology.contracts import (
    ConceptRecord,
    MappingRecord,
    MappingStatus,
    PropertyRecord,
    RelationRecord,
    SourceRecord,
)
from data_agent.ontology.package_reader import OntologyPackageReader
from data_agent.ontology.service import OntologyService

ROOT = Path(__file__).resolve().parents[1]


def _document_bytes(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_standard_parser_separates_physical_layers_and_value_domains():
    document = Document()
    document.add_paragraph("表 1 乡镇面（属性表名：a_tow_a）")
    fields = document.add_table(rows=1, cols=4)
    fields.rows[0].cells[0].text = "序号"
    fields.rows[0].cells[1].text = "字段代码"
    fields.rows[0].cells[2].text = "字段名称"
    fields.rows[0].cells[3].text = "字段类型"
    row = fields.add_row().cells
    row[0].text = "1"
    row[1].text = "BSM"
    row[2].text = "标识码"
    row[3].text = "varchar"
    row = fields.add_row().cells
    row[0].text = "2"
    row[1].text = "注：该表以县级行政区为单位填写。"
    row[2].text = "注释"
    row[3].text = ""

    document.add_paragraph("附录A 政务兴趣点分类代码表")
    document.add_paragraph("表 2 政务兴趣点分类代码表")
    codes = document.add_table(rows=1, cols=4)
    for cell, label in zip(codes.rows[0].cells, ("序号", "图层", "数据内容", "代码")):
        cell.text = label
    for ordinal, label, code in (
        ("1", "中央人民政府", "P190101"),
        ("2", "省级政府", "P190102"),
    ):
        row = codes.add_row().cells
        row[0].text = ordinal
        row[1].text = "政务兴趣点"
        row[2].text = label
        row[3].text = code

    entries, value_domains, issues = _parse_standard_document(
        _document_bytes(document), "volume-01.docx", "01"
    )

    assert [entry.code for entry in entries] == ["a_tow_a"]
    assert [field.code for field in entries[0].fields] == ["BSM"]
    assert len(value_domains) == 1
    assert [member.code for member in value_domains[0].members] == ["P190101", "P190102"]
    assert [issue["code"] for issue in issues] == [
        "invalid_standard_field_code_row_excluded"
    ]


def test_standard_parser_merges_wrapped_members_but_preserves_code_conflicts():
    document = Document()
    document.add_paragraph("附录A 分类代码表")
    document.add_paragraph("表 1 分类代码表")
    codes = document.add_table(rows=1, cols=3)
    for cell, label in zip(codes.rows[0].cells, ("序号", "数据内容", "代码")):
        cell.text = label
    for ordinal, label, code in (
        ("1", "博物馆、陈", "P1"),
        ("1", "列馆", "P1"),
        ("2", "疗养院", "P2"),
        ("3", "度假村", "P2"),
    ):
        row = codes.add_row().cells
        row[0].text = ordinal
        row[1].text = label
        row[2].text = code

    _, value_domains, issues = _parse_standard_document(
        _document_bytes(document), "volume-01.docx", "01"
    )

    members = value_domains[0].members
    assert [(member.code, member.label) for member in members] == [
        ("P1", "博物馆、陈列馆"),
        ("P2", "疗养院"),
        ("P2", "度假村"),
    ]
    assert [issue["code"] for issue in issues] == ["duplicate_value_domain_member_code"]

    compiled = compile_ontology(
        EAInput(
            packages=[], tables=[], attributes=[], connectors=[],
            source_sha256="b" * 64, source_metadata={},
        ),
        StandardInput(
            sources=[SourceRecord(
                source_id="std-doc-01",
                source_kind="standard_document",
                title="Volume 1",
                locator="test:volume-1",
                sha256="c" * 64,
            )],
            entries=[],
            value_domains=value_domains,
            issues=issues,
        ),
    )
    conflict_members = [
        concept for concept in compiled.concepts
        if concept.kind == "ValueDomainMember" and concept.code == "P2"
    ]
    assert len(conflict_members) == 2
    assert len({concept.uri for concept in conflict_members}) == 2


def _build_test_package(tmp_path, *, include_effective_fixture: bool = False):
    source = SourceRecord(
        source_id="test-source",
        source_kind="manual_governance",
        title="Test source",
        locator="test:fixture",
        sha256="a" * 64,
    )
    concepts = [
        ConceptRecord(
            concept_id="test:domain",
            uri="https://example.test/domain",
            kind="Domain",
            code="01",
            pref_label="测试领域",
            domain_id="01",
            source_system="test",
            source_id="test-source",
        ),
        ConceptRecord(
            concept_id="test:feature:a",
            uri="https://example.test/feature/a",
            kind="FeatureType",
            code="A_TABLE",
            pref_label="测试表A",
            domain_id="01",
            source_system="test",
            source_id="test-source",
        ),
        ConceptRecord(
            concept_id="test:feature:b",
            uri="https://example.test/feature/b",
            kind="FeatureType",
            code="B_TABLE",
            pref_label="测试表B",
            domain_id="01",
            source_system="test",
            source_id="test-source",
        ),
    ]
    if include_effective_fixture:
        concepts.extend([
            ConceptRecord(
                concept_id="test:class:parent",
                uri="https://example.test/class/parent",
                kind="DomainClass",
                code="ParentClass",
                pref_label="父级语义类",
                domain_id="01",
                source_system="curated_domain",
                source_id="test-source",
            ),
            ConceptRecord(
                concept_id="test:class:child",
                uri="https://example.test/class/child",
                kind="DomainClass",
                code="ChildClass",
                pref_label="子级语义类",
                domain_id="01",
                source_system="curated_domain",
                source_id="test-source",
            ),
            ConceptRecord(
                concept_id="test:schema:mapped",
                uri="https://example.test/schema/mapped",
                kind="SchemaArtifact",
                code="MAPPED_TABLE",
                pref_label="映射标准表",
                domain_id="01",
                source_system="standard",
                source_id="test-source",
            ),
        ])
    properties = [
        PropertyRecord(
            property_id="test:property:a:bsm",
            owner_concept_id="test:feature:a",
            uri="https://example.test/feature/a/bsm",
            code="BSM",
            pref_label="标识码",
            datatype="xsd:string",
            source_id="test-source",
        ),
        PropertyRecord(
            property_id="test:property:b:bsm",
            owner_concept_id="test:feature:b",
            uri="https://example.test/feature/b/bsm",
            code="BSM",
            pref_label="业务码",
            datatype="xsd:string",
            source_id="test-source",
        ),
    ]
    if include_effective_fixture:
        properties.extend([
            PropertyRecord(
                property_id="test:property:child:direct",
                owner_concept_id="test:class:child",
                uri="https://example.test/class/child/direct",
                code="directProperty",
                pref_label="直接属性",
                datatype="xsd:string",
                source_id="test-source",
            ),
            PropertyRecord(
                property_id="test:property:parent:inherited",
                owner_concept_id="test:class:parent",
                uri="https://example.test/class/parent/inherited",
                code="inheritedProperty",
                pref_label="继承属性",
                datatype="xsd:string",
                source_id="test-source",
            ),
            PropertyRecord(
                property_id="test:property:mapped:field",
                owner_concept_id="test:schema:mapped",
                uri="https://example.test/schema/mapped/field",
                code="MAPPED_FIELD",
                pref_label="映射字段",
                datatype="xsd:string",
                source_id="test-source",
            ),
        ])
    relations = [
        RelationRecord(
            relation_id="test:relation:contains:a",
            relation_type="contains",
            source_concept_id="test:domain",
            target_concept_id="test:feature:a",
            source_id="test-source",
        ),
        RelationRecord(
            relation_id="test:relation:a:b",
            relation_type="associatedWith",
            source_concept_id="test:feature:a",
            target_concept_id="test:feature:b",
            source_id="test-source",
        ),
    ]
    if include_effective_fixture:
        relations.extend([
            RelationRecord(
                relation_id="test:relation:child:parent",
                relation_type="subClassOf",
                source_concept_id="test:class:child",
                target_concept_id="test:class:parent",
                source_id="test-source",
            ),
        ])
    mappings = (
        [
            MappingRecord(
                mapping_id="test:mapping:schema:child",
                source_concept_id="test:schema:mapped",
                target_concept_id="test:class:child",
                mapping_type="describes",
                mapping_status=MappingStatus.CONFIRMED,
                confidence=1.0,
                evidence={"match_basis": ["curated_binding"]},
                reviewed_by="test-reviewer",
                reviewed_at=datetime(2026, 8, 4, tzinfo=UTC),
            ),
        ]
        if include_effective_fixture
        else []
    )
    compiled = CompiledOntology(
        sources=[source],
        concepts=concepts,
        properties=properties,
        relations=relations,
        mappings=mappings,
        issues=[],
    )
    package_dir = tmp_path / "1.0.0"
    write_package(
        compiled,
        package_dir,
        semantic_version="1.0.0",
        generated_at=datetime(2026, 8, 4, tzinfo=UTC),
    )
    return package_dir


def test_package_reader_indexes_fields_and_layouts_only_the_selected_subgraph(tmp_path):
    reader = OntologyPackageReader(_build_test_package(tmp_path))

    candidates = reader.property_candidates(
        code="bsm",
        domain_id="01",
        owner_kinds={"FeatureType"},
    )
    assert len(candidates) == 2
    assert {candidate["owner_concept_id"] for candidate in candidates} == {
        "test:feature:a",
        "test:feature:b",
    }

    graph = reader.graph(root_id="test:domain", depth=2, limit=10)
    assert graph["node_count"] == 3
    positions = {
        (node["position"]["x"], node["position"]["y"])
        for node in graph["nodes"]
    }
    assert len(positions) == graph["node_count"]


def test_package_reader_groups_direct_inherited_and_confirmed_mapped_properties(tmp_path):
    reader = OntologyPackageReader(
        _build_test_package(
            tmp_path,
            include_effective_fixture=True,
        )
    )

    direct = reader.properties("test:class:child")
    effective = reader.properties(
        "test:class:child",
        include_effective=True,
        limit=20,
    )

    assert direct["total"] == 1
    assert effective["total"] == 3
    assert effective["group_counts"] == {
        "direct": 1,
        "inherited": 1,
        "mapped": 1,
    }
    by_code = {item["code"]: item for item in effective["items"]}
    assert by_code["directProperty"]["origin_type"] == "direct"
    assert by_code["inheritedProperty"]["origin_type"] == "inherited"
    assert by_code["inheritedProperty"]["origin_depth"] == 1
    assert by_code["MAPPED_FIELD"]["origin_type"] == "mapped"
    assert by_code["MAPPED_FIELD"]["origin_concept"]["code"] == "MAPPED_TABLE"
    assert by_code["MAPPED_FIELD"]["mapping"]["mapping_status"] == "confirmed"


def test_service_reports_ambiguous_exact_field_matches(tmp_path, monkeypatch):
    package_dir = _build_test_package(tmp_path)
    monkeypatch.setenv("ONTOLOGY_RUNTIME_BACKEND", "package")
    service = OntologyService(package_dir)

    result = service.align_fields([{"code": "BSM"}], domain_id="01")

    assert result["results"][0]["resolution_status"] == "ambiguous"
    assert len(result["results"][0]["candidates"]) == 2


def test_projection_cli_defaults_to_the_governed_active_package():
    args = build_parser().parse_args([
        "project",
        "--graph-store-endpoint",
        "http://ontology-rdf:3030/ontology/data?default",
    ])

    assert args.package_dir is None


def test_curated_domain_source_kind_is_admitted_by_the_authority_schema():
    migration = (
        ROOT / "data_agent/migrations/133_ontology_curated_domain_source.sql"
    ).read_text(encoding="utf-8")

    assert "curated_domain_ontology" in migration
    assert "VALIDATE CONSTRAINT ck_gda_ontology_source_kind" in migration


def test_curated_lifecycle_and_mapping_types_are_admitted_by_the_authority_schema():
    migration = (
        ROOT / "data_agent/migrations/134_ontology_curated_model_contract.sql"
    ).read_text(encoding="utf-8")

    assert "'curated'" in migration
    for mapping_type in ("denotes_class", "describes", "schema_correspondence"):
        assert f"'{mapping_type}'" in migration
    assert "VALIDATE CONSTRAINT ck_gda_ontology_concept_status" in migration
    assert "VALIDATE CONSTRAINT ck_gda_ontology_mapping_type" in migration


def test_postgres_graph_layout_has_dedicated_curated_model_lanes():
    assert PostgresOntologyReader._kind_lane("DomainClass") == 1
    assert PostgresOntologyReader._kind_lane("StateClass") == 2
    assert PostgresOntologyReader._kind_lane("ProcessClass") == 3
    assert PostgresOntologyReader._kind_lane("SchemaArtifact") == 4


def test_postgres_validation_unwraps_release_report_evidence():
    reader = object.__new__(PostgresOntologyReader)
    reader.version_id = "00000000-0000-0000-0000-000000000001"
    reader.engine = MagicMock()
    connection = reader.engine.connect.return_value.__enter__.return_value
    connection.execute.return_value.scalar.return_value = {
        "validation": {
            "conforms": True,
            "issue_count": 2,
            "severity_counts": {"warning": 2},
        },
        "competency_report": {"conforms": True},
        "semantic_quality_report": {"conforms": True},
    }

    assert reader.validation() == {
        "conforms": True,
        "issue_count": 2,
        "severity_counts": {"warning": 2},
    }
