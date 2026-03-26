import os
import re
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Pipeline-specific default titles
PIPELINE_TITLES = {
    "optimization": "空间布局优化分析报告",
    "governance": "数据质量治理审计报告",
    "general": "空间数据分析报告",
}


def _set_cell_background(cell, fill_color):
    """Helper to set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_page_number(paragraph):
    """Insert a PAGE number field into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char_begin)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    run._r.append(instr_text)

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char_end)


def _setup_styles(doc):
    """Configure document styles for professional appearance."""
    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Heading styles — dark blue color
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Microsoft YaHei'
        h_style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        h_style._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), 'Microsoft YaHei')


def _setup_page(doc, logo_path=None):
    """Configure page layout: A4, margins, header, footer."""
    section = doc.sections[0]

    # A4 page size
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # Standard margins
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # --- Header ---
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if logo_path and os.path.exists(logo_path):
        # Insert logo on the left via a separate run
        logo_run = header_para.add_run()
        logo_run.add_picture(logo_path, width=Cm(2.0))
        header_para.add_run("    ")  # spacer

    run = header_para.add_run("GIS Data Agent 分析报告")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Header bottom border (thin gray line)
    pPr = header_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- Footer: centered page number ---
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    prefix_run = footer_para.add_run("— ")
    prefix_run.font.size = Pt(8)
    prefix_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    _add_page_number(footer_para)
    suffix_run = footer_para.add_run(" —")
    suffix_run.font.size = Pt(8)
    suffix_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_cover_page(doc, title, author, pipeline_type):
    """Add title and metadata block at the top of the report."""
    # Main title
    title_text = title or PIPELINE_TITLES.get(pipeline_type, "空间数据分析报告")
    title_p = doc.add_heading(title_text, 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata line
    meta_lines = [
        f"生成日期：{date.today().strftime('%Y年%m月%d日')}",
        f"分析师：{author}" if author else "由 GIS Data Agent 自动生成",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Separator line
    doc.add_paragraph("—" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")  # spacer


def _add_qc_cover_page(doc, title, metadata: dict = None):
    """Add a professional QC report cover page with metadata fields.

    Args:
        doc: Document object
        title: Report title
        metadata: Dict with optional keys: project_name, org_name, check_org,
                  product_type, standard_id, check_date, checker, reviewer
    """
    meta = metadata or {}

    # Main title
    title_p = doc.add_heading(title or "测绘成果质量检查报告", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    # Metadata table (2-column layout)
    fields = [
        ("项目名称", meta.get("project_name", "")),
        ("委托单位", meta.get("org_name", "")),
        ("检查单位", meta.get("check_org", "")),
        ("成果类型", meta.get("product_type", "")),
        ("检查依据", meta.get("standard_id", "GB/T 24356-2009")),
        ("检查日期", meta.get("check_date", date.today().strftime("%Y年%m月%d日"))),
        ("检查人员", meta.get("checker", "")),
        ("审核人员", meta.get("reviewer", "")),
    ]

    # Filter out empty entries but keep at least check_date and standard
    fields = [(k, v) for k, v in fields if v]

    if fields:
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(fields):
            cell_label = table.cell(i, 0)
            cell_value = table.cell(i, 1)
            cell_label.text = label
            cell_value.text = value
            # Bold labels
            for para in cell_label.paragraphs:
                for run in para.runs:
                    run.font.bold = True
            # Set label column width
            cell_label.width = Cm(4)
            _set_cell_background(cell_label, "F2F2F2")

    doc.add_paragraph("")  # spacer
    doc.add_page_break()


def _add_toc(doc):
    """Insert a Table of Contents field (auto-updated on open in Word)."""
    toc_heading = doc.add_heading("目  录", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # TOC field: begin
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    # TOC instruction
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr_text)

    # TOC field: separate
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)

    # Placeholder text
    placeholder_run = paragraph.add_run("（请在 Word 中按 F9 更新目录）")
    placeholder_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # TOC field: end
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)

    doc.add_page_break()


def _fill_table_rows(doc, table_index, headers, rows):
    """Fill or create a table with headers and data rows.

    Args:
        doc: Document object
        table_index: Index of existing table (-1 to create new)
        headers: List of column header strings
        rows: List of row dicts or lists
    """
    if table_index >= 0 and table_index < len(doc.tables):
        table = doc.tables[table_index]
    else:
        # Create new table
        num_cols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.style = "Table Grid"

    # Header row
    if table.rows:
        for j, h in enumerate(headers):
            if j < len(table.rows[0].cells):
                cell = table.rows[0].cells[j]
                cell.text = h
                _set_cell_background(cell, "1F497D")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
                        run.font.size = Pt(9)

    # Data rows
    for i, row_data in enumerate(rows):
        row_idx = i + 1
        if row_idx >= len(table.rows):
            table.add_row()
        for j, val in enumerate(row_data if isinstance(row_data, (list, tuple)) else row_data.values()):
            if j < len(table.rows[row_idx].cells):
                table.rows[row_idx].cells[j].text = str(val)

    return table


def _insert_image(doc, image_path, width_cm=15, caption=""):
    """Insert an image with optional caption.

    Args:
        doc: Document object
        image_path: Path to image file
        width_cm: Image width in cm
        caption: Optional caption text below image
    """
    if not os.path.exists(image_path):
        doc.add_paragraph(f"[图片未找到: {image_path}]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))

    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _embed_chart(doc, chart_data, title="", width_cm=14):
    """Generate a matplotlib chart and embed it as an image.

    Args:
        doc: Document object
        chart_data: Dict with 'type' (bar|pie|line), 'labels', 'values', optional 'colors'
        title: Chart title
        width_cm: Image width in cm
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import tempfile

        fig, ax = plt.subplots(figsize=(8, 5))

        chart_type = chart_data.get("type", "bar")
        labels = chart_data.get("labels", [])
        values = chart_data.get("values", [])
        colors = chart_data.get("colors", None)

        if chart_type == "bar":
            ax.bar(labels, values, color=colors or "#1F497D")
        elif chart_type == "pie":
            ax.pie(values, labels=labels, colors=colors, autopct="%1.1f%%", startangle=90)
        elif chart_type == "line":
            ax.plot(labels, values, marker="o", color=colors[0] if colors else "#1F497D")
        else:
            ax.bar(labels, values, color="#1F497D")

        if title:
            ax.set_title(title, fontsize=12)
        plt.tight_layout()

        # Save to temp file and insert
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
            fig.savefig(tmp_path, dpi=150, bbox_inches="tight")
        plt.close(fig)

        _insert_image(doc, tmp_path, width_cm=width_cm, caption=title)

        # Cleanup
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
    except ImportError:
        doc.add_paragraph(f"[图表生成失败: matplotlib 未安装] {title}")
    except Exception as e:
        doc.add_paragraph(f"[图表生成失败: {e}]")


def generate_qc_report(
    section_data: dict,
    metadata: dict = None,
    charts: list = None,
    images: list = None,
    output_dir: str = None,
) -> str:
    """Generate a professional surveying QC report with cover, TOC, tables, and charts.

    Args:
        section_data: Dict mapping section names to markdown content or table data.
            Keys should match QC template sections:
            "项目概况", "检查依据", "数据审查结果", "精度核验结果",
            "缺陷统计", "质量评分", "整改建议", "结论"
        metadata: Cover page metadata dict (project_name, org_name, etc.)
        charts: List of chart dicts to embed (type, labels, values, title)
        images: List of image dicts (path, caption, after_section)
        output_dir: Output directory (defaults to user upload dir)

    Returns:
        Absolute path to the generated .docx file.
    """
    import uuid as _uuid

    doc = Document()
    _setup_styles(doc)
    _setup_page(doc)

    # 1. Cover page
    title = (metadata or {}).get("title", "测绘成果质量检查报告")
    _add_qc_cover_page(doc, title, metadata)

    # 2. Table of contents
    _add_toc(doc)

    # 3. Sections
    sections = [
        "项目概况", "检查依据", "数据审查结果", "精度核验结果",
        "缺陷统计", "质量评分", "整改建议", "结论",
    ]

    for section_name in sections:
        content = section_data.get(section_name, "")
        doc.add_heading(section_name, level=1)

        if isinstance(content, dict) and "table" in content:
            # Table data: {"table": {"headers": [...], "rows": [...]}}
            tdata = content["table"]
            _fill_table_rows(doc, -1, tdata.get("headers", []), tdata.get("rows", []))
            if content.get("text"):
                doc.add_paragraph("")
                _render_markdown_body(doc, content["text"])
        elif content:
            _render_markdown_body(doc, content)
        else:
            doc.add_paragraph(f"（{section_name}暂无数据）")

        # Insert images for this section
        if images:
            for img in images:
                if img.get("after_section") == section_name:
                    _insert_image(doc, img["path"], caption=img.get("caption", ""))

    # 4. Embed charts at the end (or per-section if specified)
    if charts:
        for chart in charts:
            if not chart.get("after_section"):
                _embed_chart(doc, chart, title=chart.get("title", ""))

    # 5. Save
    uid = _uuid.uuid4().hex[:8]
    if not output_dir:
        from .user_context import get_user_upload_dir
        output_dir = get_user_upload_dir()
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"qc_report_{uid}.docx")
    doc.save(path)
    return os.path.abspath(path)


def _render_markdown_body(doc, markdown_text):
    """Parse markdown text and render into the Word document."""
    lines = markdown_text.split('\n')
    img_pattern = r'[^<>:"|?*\s]+\.png'

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # --- 1. Table Detection & Handling ---
        if line.startswith('|') and i + 1 < len(lines) and re.match(r'^[|\s\-:]+$', lines[i + 1].strip()):
            table_buffer = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not re.match(r'^[|\s\-:]+$', lines[i].strip()):
                    row_data = [cell.strip() for cell in lines[i].strip('|').split('|')]
                    table_buffer.append(row_data)
                i += 1

            if table_buffer:
                rows = len(table_buffer)
                cols = max(len(r) for r in table_buffer)
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'

                for r_idx, row_cells in enumerate(table_buffer):
                    for c_idx, val in enumerate(row_cells):
                        if c_idx < cols:
                            cell = table.cell(r_idx, c_idx)
                            p = cell.paragraphs[0]
                            parts = re.split(r'(\*\*.*?\*\*)', val)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                else:
                                    p.add_run(part)

                            if r_idx == 0:
                                _set_cell_background(cell, "4472C4")
                                if p.runs:
                                    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                                    p.runs[0].bold = True
                doc.add_paragraph()
            continue

        if not line:
            i += 1
            continue

        # --- 2. Image Handling ---
        img_match = re.search(img_pattern, line, re.IGNORECASE)
        if img_match:
            img_path = img_match.group(0)
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = os.path.basename(img_path)
                    p = doc.add_paragraph(f"图：{caption}")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                except Exception:
                    doc.add_paragraph(f"[图片加载失败: {img_path}]")
            i += 1
            continue

        # --- 3. Headers ---
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)

        # --- 4. Lists ---
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(line, style='List Number')

        # --- 5. Text with bold ---
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1


def generate_word_report(
    markdown_text: str,
    output_path: str = "report.docx",
    title: str = None,
    author: str = None,
    logo_path: str = None,
    pipeline_type: str = "general",
) -> str:
    """
    Convert Markdown text to a professionally formatted Word document.

    Args:
        markdown_text: Markdown-formatted analysis text (with tables, images, headers).
        output_path: Output .docx file path.
        title: Custom report title. If None, auto-selected by pipeline_type.
        author: Author name (from user display_name).
        logo_path: Optional path to a logo image for the header.
        pipeline_type: Pipeline type for default title selection.
    Returns:
        Absolute path to the generated .docx file.
    """
    doc = Document()

    _setup_styles(doc)
    _setup_page(doc, logo_path)
    _add_cover_page(doc, title, author, pipeline_type)
    _render_markdown_body(doc, markdown_text)

    doc.save(output_path)
    return os.path.abspath(output_path)


def generate_pdf_report(
    markdown_text: str,
    output_path: str = "report.pdf",
    title: str = None,
    author: str = None,
    logo_path: str = None,
    pipeline_type: str = "general",
) -> str:
    """
    Generate a PDF report by first creating a Word document, then converting to PDF.
    Falls back to returning the Word document if PDF conversion is unavailable.

    Args:
        markdown_text: Markdown-formatted analysis text.
        output_path: Output .pdf file path.
        title: Custom report title.
        author: Author name.
        logo_path: Optional logo image path.
        pipeline_type: Pipeline type for default title selection.
    Returns:
        Absolute path to the generated file (.pdf or .docx fallback).
    """
    # Generate Word document first
    docx_path = output_path.rsplit('.', 1)[0] + '.docx'
    generate_word_report(markdown_text, docx_path, title, author, logo_path, pipeline_type)

    # Attempt PDF conversion
    try:
        from docx2pdf import convert
        convert(docx_path, output_path)
        # Clean up temporary Word file
        try:
            os.remove(docx_path)
        except OSError:
            pass
        return os.path.abspath(output_path)
    except Exception as e:
        print(f"[Report] PDF conversion failed ({e}). Returning Word document instead.")
        return os.path.abspath(docx_path)


# =====================================================================
# Template-based Report Generation (v15.1 — 测绘质检报告扩展)
# =====================================================================

# Built-in report templates
REPORT_TEMPLATES = {
    "surveying_qc": {
        "id": "surveying_qc",
        "name": "测绘质检报告",
        "description": "测绘成果质量检查与验收报告（符合 GB/T 24356）",
        "pipeline_type": "governance",
        "sections": [
            "项目概况", "检查依据", "数据审查结果",
            "精度核验结果", "缺陷统计", "质量评分", "整改建议", "结论",
        ],
    },
    "data_quality": {
        "id": "data_quality",
        "name": "数据质量报告",
        "description": "空间数据质量评估报告",
        "pipeline_type": "governance",
        "sections": [
            "数据集概览", "质量评估", "问题清单", "改进建议",
        ],
    },
    "governance": {
        "id": "governance",
        "name": "数据治理报告",
        "description": "数据治理综合评估报告",
        "pipeline_type": "governance",
        "sections": [
            "治理概览", "标准符合性", "质量评分", "Gap分析", "治理建议",
        ],
    },
    "general_analysis": {
        "id": "general_analysis",
        "name": "空间分析报告",
        "description": "通用空间数据分析报告",
        "pipeline_type": "general",
        "sections": [
            "分析概览", "数据描述", "分析结果", "可视化", "结论",
        ],
    },
}


def list_report_templates() -> list[dict]:
    """List available report templates."""
    return [
        {"id": t["id"], "name": t["name"], "description": t["description"]}
        for t in REPORT_TEMPLATES.values()
    ]


def generate_structured_report(
    template_id: str,
    section_data: dict[str, str],
    title: str = None,
    author: str = None,
    output_format: str = "docx",
    output_dir: str = None,
) -> str:
    """Generate a report from a template with section data.

    Args:
        template_id: Template identifier (e.g., 'surveying_qc').
        section_data: Dict mapping section names to markdown content.
        title: Report title override.
        author: Author name.
        output_format: 'docx', 'pdf', or 'md'.
        output_dir: Output directory.

    Returns:
        Absolute path to the generated report file.
    """
    template = REPORT_TEMPLATES.get(template_id)
    if not template:
        available = ", ".join(REPORT_TEMPLATES.keys())
        raise ValueError(f"Unknown template: {template_id}. Available: {available}")

    # Build markdown from sections
    report_title = title or template["name"]
    lines = [f"# {report_title}\n"]

    for section_name in template["sections"]:
        content = section_data.get(section_name, "")
        lines.append(f"## {section_name}\n")
        if content:
            lines.append(content + "\n")
        else:
            lines.append(f"*（{section_name}暂无数据）*\n")

    markdown_text = "\n".join(lines)

    # Determine output path
    import uuid as _uuid
    uid = _uuid.uuid4().hex[:8]
    if not output_dir:
        from .user_context import get_user_upload_dir
        output_dir = get_user_upload_dir()
    os.makedirs(output_dir, exist_ok=True)

    if output_format == "md":
        path = os.path.join(output_dir, f"report_{uid}.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write(markdown_text)
        return os.path.abspath(path)
    elif output_format == "pdf":
        path = os.path.join(output_dir, f"report_{uid}.pdf")
        return generate_pdf_report(
            markdown_text, path, report_title, author,
            pipeline_type=template.get("pipeline_type", "general"),
        )
    else:
        path = os.path.join(output_dir, f"report_{uid}.docx")
        return generate_word_report(
            markdown_text, path, report_title, author,
            pipeline_type=template.get("pipeline_type", "general"),
        )
