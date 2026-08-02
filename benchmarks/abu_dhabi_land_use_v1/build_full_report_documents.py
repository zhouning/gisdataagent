#!/usr/bin/env python3
"""Build polished DOCX and PDF deliverables from the full Markdown report."""

from __future__ import annotations

import argparse
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
STEM = "阿布扎比土地利用三模型模拟与优化完整报告_2026-08-02"
DEFAULT_MARKDOWN = HERE / f"{STEM}.md"
DEFAULT_DOCX = HERE / f"{STEM}.docx"
DEFAULT_PDF = HERE / f"{STEM}.pdf"
PDF_HEADER = HERE / "report_pdf_header.tex"


def _set_font(style: object, *, latin: str, east_asia: str, size: float) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def _set_run_font(run: object, *, size: float | None = None) -> None:
    run.font.name = "Arial"
    if size is not None:
        run.font.size = Pt(size)
    properties = run._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:cs"), "Arial")
    fonts.set(qn("w:eastAsia"), "Arial Unicode MS")


def _page_number(paragraph: object) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def _repeat_table_header(row: object) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def polish_docx(path: Path) -> None:
    document = Document(path)
    normal = document.styles["Normal"]
    _set_font(normal, latin="Arial", east_asia="Arial Unicode MS", size=10.5)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.space_after = Pt(5)

    heading_settings = {
        "Title": (22, "1F3A4D"),
        "Subtitle": (12, "52636E"),
        "Heading 1": (16, "1F3A4D"),
        "Heading 2": (13, "285C6A"),
        "Heading 3": (11.5, "3D6D73"),
    }
    for name, (size, color) in heading_settings.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        _set_font(style, latin="Arial", east_asia="Arial Unicode MS", size=size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "Caption" in document.styles:
        caption = document.styles["Caption"]
        _set_font(caption, latin="Arial", east_asia="Arial Unicode MS", size=9)
        caption.font.color.rgb = RGBColor(0x4B, 0x55, 0x5B)
        caption.font.italic = False
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(8)

    for section_index, section in enumerate(document.sections):
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
        section.different_first_page_header_footer = True
        header = section.header.paragraphs[0]
        header.text = "阿布扎比土地利用三模型模拟与优化完整报告"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            _set_run_font(run, size=8)
            run.font.color.rgb = RGBColor(0x6D, 0x77, 0x7C)
        footer = section.footer.paragraphs[0]
        _page_number(footer)
        if section_index > 0 and section.start_type == WD_SECTION.NEW_PAGE:
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
        if section_index == 0:
            section.first_page_header.paragraphs[0].clear()
            section.first_page_footer.paragraphs[0].clear()

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run)
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(4)
        if paragraph.style.name == "Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(120)
            paragraph.paragraph_format.space_after = Pt(18)
        elif paragraph.style.name in {"Subtitle", "Author", "Date"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.style.name == "Subtitle":
            paragraph.paragraph_format.space_after = Pt(48)
        if paragraph.style.name == "Date" and paragraph.text.strip() == "2026-08-02":
            paragraph.add_run().add_break(WD_BREAK.PAGE)
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "执行摘要":
            paragraph.paragraph_format.page_break_before = True

    available_styles = {style.name for style in document.styles}
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for candidate in ("Light Shading Accent 1", "Table Grid"):
            if candidate in available_styles:
                table.style = candidate
                break
        if table.rows:
            _repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(1.5)
                    paragraph.paragraph_format.line_spacing = 1.05
                    for run in paragraph.runs:
                        _set_run_font(
                            run,
                            size=8 if len(table.columns) >= 7 else 8.5,
                        )
                        if row_index == 0:
                            run.font.bold = True

    document.core_properties.title = "阿布扎比土地利用三模型模拟与优化完整报告"
    document.core_properties.subject = "GeoSOS-FLUS、Geospatial Kernel 与 Paper58 统一实验"
    document.core_properties.author = "GWM / Geospatial Kernel 研究工作"
    document.save(path)


def build(*, markdown: Path, docx: Path, pdf: Path) -> None:
    subprocess.run(
        [
            "pandoc",
            str(markdown),
            "--from=markdown+yaml_metadata_block+implicit_figures+tex_math_dollars+pipe_tables+link_attributes",
            "--standalone",
            f"--resource-path={HERE}",
            "--reference-location=block",
            "-o",
            str(docx),
        ],
        cwd=HERE,
        check=True,
    )
    polish_docx(docx)
    subprocess.run(
        [
            "pandoc",
            str(markdown),
            "--from=markdown+yaml_metadata_block+implicit_figures+tex_math_dollars+pipe_tables+link_attributes",
            "--standalone",
            "--pdf-engine=xelatex",
            f"--resource-path={HERE}",
            f"--include-in-header={PDF_HEADER}",
            "--variable=mainfont:Arial Unicode MS",
            "--variable=monofont:Arial Unicode MS",
            "--variable=papersize:a4",
            "--variable=fontsize:10pt",
            "--variable=geometry:top=22mm,bottom=20mm,left=20mm,right=20mm,headheight=15pt",
            "--variable=figure-placement:H",
            "--variable=colorlinks:true",
            "--variable=linkcolor:teal",
            "--variable=urlcolor:teal",
            "-o",
            str(pdf),
        ],
        cwd=HERE,
        check=True,
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", type=Path, default=DEFAULT_MARKDOWN)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--pdf", type=Path, default=DEFAULT_PDF)
    args = parser.parse_args()
    build(markdown=args.markdown, docx=args.docx, pdf=args.pdf)
    print(args.docx)
    print(args.pdf)


if __name__ == "__main__":
    main()
